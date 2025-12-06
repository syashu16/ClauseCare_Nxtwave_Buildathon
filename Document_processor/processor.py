"""
Advanced Document Processor for ClauseCare.
Handles PDF, DOCX, TXT, and Images.
Features: OCR, Metadata Extraction, Text Chunking, and Content Analysis.

Note: OCR functionality requires Tesseract installation (optional for cloud deployment).
"""

import io
import re
from datetime import datetime
from typing import Dict, Any, List, Union, Tuple

# PDF processing - PyMuPDF
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

# DOCX processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = None

# Image processing
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    Image = None

# OCR - requires system Tesseract (optional)
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    pytesseract = None

# Text analysis
try:
    import textstat
    TEXTSTAT_AVAILABLE = True
except ImportError:
    TEXTSTAT_AVAILABLE = False
    textstat = None

try:
    from langdetect import detect
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False
    detect = None

class DocumentProcessor:
    """
    Enterprise-grade document processor for extracting, cleaning, and analyzing text.
    Works on Streamlit Cloud without requiring Tesseract OCR.
    """

    def __init__(self, tesseract_cmd: str = None):
        """
        Initialize processor. 
        Args:
            tesseract_cmd: Optional path to tesseract executable if not in PATH.
        """
        if tesseract_cmd and TESSERACT_AVAILABLE and pytesseract:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    def process_document(self, file) -> Dict[str, Any]:
        """
        Main pipeline entry point.
        Returns a dictionary containing text, metadata, tables, stats, and analysis.
        """
        try:
            filename = file.name.lower() if hasattr(file, 'name') else "unknown"
            file_ext = filename.split('.')[-1]

            # 1. Extraction Strategy
            tables = []
            if file_ext == 'pdf':
                if not PYMUPDF_AVAILABLE:
                    raise ImportError("PyMuPDF not installed. Install with: pip install pymupdf")
                text, metadata, tables = self._read_pdf(file)
            elif file_ext == 'docx':
                if not DOCX_AVAILABLE:
                    raise ImportError("python-docx not installed. Install with: pip install python-docx")
                text, metadata, tables = self._read_docx(file)
            elif file_ext == 'txt':
                text, metadata = self._read_txt(file)
            elif file_ext in ['jpg', 'jpeg', 'png']:
                if not TESSERACT_AVAILABLE or not PIL_AVAILABLE:
                    raise ImportError("OCR requires Tesseract and Pillow. Image processing not available on cloud.")
                text, metadata = self._read_image(file)
            else:
                raise ValueError(f"Unsupported format: {file_ext}")

            # 2. Run Advanced Analysis (Language & Readability)
            analysis = self._analyze_text_quality(text)

            # 3. Return unified structure
            return {
                "filename": filename,
                "text": text,
                "metadata": metadata,
                "tables": tables,
                "analysis": analysis,
                "status": "success"
            }

        except Exception as e:
            return {"status": "error", "error_message": str(e)}

    # --- Analysis Method ---

    def _analyze_text_quality(self, text: str) -> Dict[str, Any]:
        """
        Analyzes the complexity and language of the text.
        """
        if not text.strip():
            return {"language": "Unknown", "readability_score": 0, "difficulty": "N/A"}

        # 1. Detect Language
        lang = "Unknown"
        if LANGDETECT_AVAILABLE and detect:
            try:
                lang = detect(text)
            except:
                lang = "Unknown"

        # 2. Calculate Readability (Flesch Reading Ease)
        score = 50  # Default score
        if TEXTSTAT_AVAILABLE and textstat:
            try:
                score = textstat.flesch_reading_ease(text)
            except:
                score = 50
        
        # 3. Interpret Score
        difficulty = "Standard"
        if score > 80: difficulty = "Very Easy (5th Grade)"
        elif score > 60: difficulty = "Easy (8th Grade)"
        elif score < 30: difficulty = "Very Complex (Academic/Legal)"
        else: difficulty = "Hard (College Level)"

        return {
            "language": lang.upper() if isinstance(lang, str) else "Unknown",
            "readability_score": score,
            "difficulty": difficulty
        }

    # --- File Handlers ---

    def _read_pdf(self, file) -> Tuple[str, Dict, List[Dict]]:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        full_text = []
        all_tables = []
        
        # Extract Metadata
        metadata = {
            "author": doc.metadata.get('author', 'unknown') if doc.metadata else 'unknown',
            "creation_date": doc.metadata.get('creationDate', 'unknown') if doc.metadata else 'unknown',
            "pages": doc.page_count,
            "type": "PDF"
        }

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            # Fallback to OCR if page seems empty (scanned) - only if Tesseract available
            if len(text.strip()) < 50 and TESSERACT_AVAILABLE and PIL_AVAILABLE and pytesseract and Image:
                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = pix.tobytes("png")
                    text = pytesseract.image_to_string(Image.open(io.BytesIO(img_data)))
                except:
                    # OCR failed, use whatever text we got
                    pass
            
            full_text.append(text)
            
            # Extract tables from page
            tables = self._extract_tables_from_pdf_page(page, page_num)
            all_tables.extend(tables)
            
        doc.close()
        return "\n".join(full_text), metadata, all_tables

    def _read_docx(self, file) -> Tuple[str, Dict, List[Dict]]:
        doc = docx.Document(file)
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Extract tables from DOCX
        tables = self._extract_tables_from_docx(doc)
        
        metadata = {
            "author": doc.core_properties.author or "unknown",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "unknown",
            "paragraphs": len(doc.paragraphs),
            "tables": len(tables),
            "type": "DOCX"
        }
        return text, metadata, tables

    def _read_txt(self, file) -> Tuple[str, Dict]:
        content = file.read()
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="ignore")
        metadata = {
            "size_bytes": len(content),
            "type": "TXT"
        }
        return content, metadata

    def _read_image(self, file) -> Tuple[str, Dict]:
        image = Image.open(file)
        text = pytesseract.image_to_string(image)
        metadata = {
            "format": image.format,
            "size": image.size,
            "mode": image.mode,
            "type": "IMAGE"
        }
        return text, metadata

    # --- Text Operations ---

    def clean_text(self, text: str) -> str:
        """Standardizes text formatting."""
        if not text: return ""
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text) # Fix hyphenation
        text = re.sub(r'\n\s*\n+', '\n\n', text) # Fix excessive newlines
        return text.strip()

    def redact_pii(self, text: str) -> str:
        """Redacts Emails and Phone numbers (Basic Regex)."""
        # Email Regex
        text = re.sub(r'[\w\.-]+@[\w\.-]+\.\w+', '[REDACTED_EMAIL]', text)
        # Phone Regex (simple international format)
        text = re.sub(r'\b\+?[\d\s-]{10,}\b', '[REDACTED_PHONE]', text)
        return text

    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """
        Splits text into chunks. 
        Critical for RAG/LLM applications.
        """
        if not text: return []
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            chunks.append(chunk)
            
        return chunks
    
    def _extract_tables_from_pdf_page(self, page, page_num: int) -> List[Dict]:
        """
        Extracts tables from a PDF page using PyMuPDF.
        Returns a list of tables with their data and metadata.
        """
        tables = []
        try:
            # Find tables using pymupdf's table detection
            tabs = page.find_tables()
            
            for table_idx, table in enumerate(tabs.tables, 1):
                # Extract table data
                table_data = table.extract()
                
                if table_data:
                    # Convert to structured format
                    structured_table = {
                        "page": page_num,
                        "table_number": table_idx,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0,
                        "data": table_data,
                        "text_representation": self._format_table_as_text(table_data)
                    }
                    tables.append(structured_table)
        except Exception as e:
            # If table extraction fails, return empty list
            pass
        
        return tables
    
    def _extract_tables_from_docx(self, doc) -> List[Dict]:
        """
        Extracts all tables from a DOCX document.
        Returns a list of tables with their data and metadata.
        """
        tables = []
        
        for table_idx, table in enumerate(doc.tables, 1):
            table_data = []
            
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                structured_table = {
                    "table_number": table_idx,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                    "data": table_data,
                    "text_representation": self._format_table_as_text(table_data)
                }
                tables.append(structured_table)
        
        return tables
    
    def _format_table_as_text(self, table_data: List[List[str]]) -> str:
        """
        Formats table data as readable text for processing.
        """
        if not table_data:
            return ""
        
        # Calculate column widths
        col_widths = []
        for col_idx in range(len(table_data[0])):
            max_width = max(len(str(row[col_idx])) if col_idx < len(row) else 0 for row in table_data)
            col_widths.append(max(max_width, 10))
        
        # Format rows
        formatted_rows = []
        for row in table_data:
            formatted_cells = []
            for col_idx, cell in enumerate(row):
                if col_idx < len(col_widths):
                    formatted_cells.append(str(cell).ljust(col_widths[col_idx]))
            formatted_rows.append(" | ".join(formatted_cells))
        
        return "\n".join(formatted_rows)