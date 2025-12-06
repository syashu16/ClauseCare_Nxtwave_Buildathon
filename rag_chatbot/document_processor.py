"""
Document Processor for RAG Chatbot

Handles:
- Document loading (PDF, DOCX, TXT)
- Smart chunking with overlap
- Metadata extraction
- Section detection
"""

import re
import io
from dataclasses import dataclass, field
from typing import Optional
from datetime import datetime


@dataclass
class DocumentChunk:
    """Represents a chunk of a document"""
    chunk_id: str
    content: str
    metadata: dict = field(default_factory=dict)
    section: Optional[str] = None
    page_number: Optional[int] = None
    start_char: int = 0
    end_char: int = 0
    
    def __post_init__(self):
        if not self.metadata:
            self.metadata = {}
        self.metadata.update({
            "chunk_id": self.chunk_id,
            "section": self.section,
            "page_number": self.page_number,
        })


@dataclass  
class ProcessedDocument:
    """Represents a fully processed document"""
    doc_id: str
    filename: str
    chunks: list[DocumentChunk]
    total_pages: int = 0
    total_chars: int = 0
    sections: list[str] = field(default_factory=list)
    processed_at: datetime = field(default_factory=datetime.now)
    
    @property
    def total_chunks(self) -> int:
        return len(self.chunks)


class DocumentProcessor:
    """
    Processes legal documents into chunks for RAG
    
    Features:
    - Smart chunking that respects sentence boundaries
    - Section detection for legal documents
    - Overlap for context preservation
    - Metadata extraction
    """
    
    # Section patterns for legal documents
    SECTION_PATTERNS = [
        r'^(?:ARTICLE|Article|SECTION|Section)\s+[IVX\d]+[.:]\s*(.+)$',
        r'^(?:\d+\.)\s+([A-Z][A-Z\s]+)$',
        r'^(?:\d+\.)\s+([A-Z][a-zA-Z\s]+)$',
        r'^([A-Z][A-Z\s]{3,})$',
        r'^\s*(?:\d+\.)+\s*([A-Z][a-zA-Z\s]+)',
    ]
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        min_chunk_size: int = 100,
    ):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Target size for each chunk in characters
            chunk_overlap: Overlap between chunks for context
            min_chunk_size: Minimum chunk size to keep
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
    
    def process_file(self, file_bytes: bytes, filename: str) -> ProcessedDocument:
        """
        Process an uploaded file into chunks.
        
        Args:
            file_bytes: Raw file content
            filename: Original filename
            
        Returns:
            ProcessedDocument with chunks
        """
        # Determine file type and extract text
        if filename.lower().endswith('.pdf'):
            text, pages = self._extract_pdf(file_bytes)
        elif filename.lower().endswith('.docx'):
            text, pages = self._extract_docx(file_bytes)
        elif filename.lower().endswith('.txt'):
            text = file_bytes.decode('utf-8')
            pages = 1
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Generate document ID
        doc_id = f"doc_{hash(filename + str(len(text)))}"
        
        # Detect sections
        sections = self._detect_sections(text)
        
        # Create chunks
        chunks = self._create_chunks(text, doc_id, sections)
        
        return ProcessedDocument(
            doc_id=doc_id,
            filename=filename,
            chunks=chunks,
            total_pages=pages,
            total_chars=len(text),
            sections=list(sections.keys()),
        )
    
    def process_text(self, text: str, doc_name: str = "document") -> ProcessedDocument:
        """
        Process raw text into chunks.
        
        Args:
            text: Raw document text
            doc_name: Name for the document
            
        Returns:
            ProcessedDocument with chunks
        """
        doc_id = f"doc_{hash(doc_name + str(len(text)))}"
        sections = self._detect_sections(text)
        chunks = self._create_chunks(text, doc_id, sections)
        
        return ProcessedDocument(
            doc_id=doc_id,
            filename=doc_name,
            chunks=chunks,
            total_chars=len(text),
            sections=list(sections.keys()),
        )
    
    def _extract_pdf(self, file_bytes: bytes) -> tuple[str, int]:
        """Extract text from PDF using PyMuPDF"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text:
                    # Add page marker for citation
                    text_parts.append(f"[PAGE {page_num + 1}]\n{page_text}")
            
            num_pages = len(doc)
            doc.close()
            
            return "\n\n".join(text_parts), num_pages
            
        except ImportError:
            raise ImportError("PyMuPDF required. Install with: pip install pymupdf")
    
    def _extract_docx(self, file_bytes: bytes) -> tuple[str, int]:
        """Extract text from DOCX"""
        try:
            import docx
            
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            return "\n\n".join(paragraphs), 1
            
        except ImportError:
            raise ImportError("python-docx required. Install with: pip install python-docx")
    
    def _detect_sections(self, text: str) -> dict[str, int]:
        """
        Detect section headings and their positions.
        
        Returns:
            Dict mapping section name to character position
        """
        sections = {}
        lines = text.split('\n')
        char_pos = 0
        
        for line in lines:
            line_stripped = line.strip()
            
            for pattern in self.SECTION_PATTERNS:
                match = re.match(pattern, line_stripped)
                if match:
                    section_name = match.group(1) if match.groups() else line_stripped
                    sections[section_name.strip()] = char_pos
                    break
            
            char_pos += len(line) + 1
        
        return sections
    
    def _find_section_for_position(self, pos: int, sections: dict[str, int]) -> Optional[str]:
        """Find which section a character position belongs to"""
        current_section = None
        
        for section_name, section_pos in sorted(sections.items(), key=lambda x: x[1]):
            if section_pos <= pos:
                current_section = section_name
            else:
                break
        
        return current_section
    
    def _find_page_for_position(self, text: str, pos: int) -> Optional[int]:
        """Find which page a position is on based on [PAGE X] markers"""
        page_pattern = r'\[PAGE (\d+)\]'
        
        # Find all page markers before this position
        text_before = text[:pos]
        matches = list(re.finditer(page_pattern, text_before))
        
        if matches:
            return int(matches[-1].group(1))
        return 1
    
    def _create_chunks(
        self,
        text: str,
        doc_id: str,
        sections: dict[str, int]
    ) -> list[DocumentChunk]:
        """
        Create overlapping chunks from text.
        
        Uses sentence-aware splitting to avoid breaking mid-sentence.
        """
        chunks = []
        
        # Clean text but preserve structure
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Split into sentences for better chunking
        sentences = self._split_into_sentences(text)
        
        current_chunk = []
        current_length = 0
        chunk_start = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_len = len(sentence)
            
            # If adding this sentence exceeds chunk size
            if current_length + sentence_len > self.chunk_size and current_chunk:
                # Create chunk
                chunk_text = ' '.join(current_chunk)
                chunk_end = chunk_start + len(chunk_text)
                
                if len(chunk_text) >= self.min_chunk_size:
                    chunk = DocumentChunk(
                        chunk_id=f"{doc_id}_chunk_{chunk_index}",
                        content=chunk_text,
                        section=self._find_section_for_position(chunk_start, sections),
                        page_number=self._find_page_for_position(text, chunk_start),
                        start_char=chunk_start,
                        end_char=chunk_end,
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Keep overlap sentences
                overlap_sentences = []
                overlap_length = 0
                for s in reversed(current_chunk):
                    if overlap_length + len(s) <= self.chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_length += len(s)
                    else:
                        break
                
                current_chunk = overlap_sentences
                current_length = overlap_length
                chunk_start = chunk_end - overlap_length
            
            current_chunk.append(sentence)
            current_length += sentence_len
        
        # Don't forget the last chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text) >= self.min_chunk_size:
                chunk = DocumentChunk(
                    chunk_id=f"{doc_id}_chunk_{chunk_index}",
                    content=chunk_text,
                    section=self._find_section_for_position(chunk_start, sections),
                    page_number=self._find_page_for_position(text, chunk_start),
                    start_char=chunk_start,
                    end_char=chunk_start + len(chunk_text),
                )
                chunks.append(chunk)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> list[str]:
        """Split text into sentences while preserving legal numbering"""
        # Pattern for sentence endings (avoiding legal numbering like "1." "2.")
        sentence_pattern = r'(?<=[.!?])\s+(?=[A-Z])'
        
        # Split but keep legal structure
        raw_sentences = re.split(sentence_pattern, text)
        
        # Clean sentences
        sentences = []
        for s in raw_sentences:
            s = s.strip()
            if s:
                sentences.append(s)
        
        return sentences
