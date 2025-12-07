"""
ClauseCare - Unified Legal Document Platform

Integrates:
1. Risk Assessment - Analyze contract risks with AI
2. Document Chat - RAG-powered Q&A with citations  
3. Document Processor - OCR, tables, metadata extraction
4. Clause Simplification - Plain English translations
5. NegotiateAI - Multi-Agent Negotiation Intelligence

All features work independently and share document context.
"""

import streamlit as st
import os
import io
import json
import time
from datetime import datetime
from typing import Optional
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Get API key from environment
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")

# Import Risk Assessment components
try:
    from risk_assessment import (
        RiskAssessmentEngine,
        RiskCategory,
        SeverityLevel,
        DocumentRisk,
    )
    from risk_assessment.ai_analyzer import AnalysisContext
    from risk_assessment.risk_assessment_engine import ProgressCallback
    RISK_ASSESSMENT_AVAILABLE = True
except ImportError as e:
    RISK_ASSESSMENT_AVAILABLE = False
    print(f"Risk assessment not available: {e}")

# Import RAG Chatbot components
try:
    from rag_chatbot.document_processor import DocumentProcessor as RAGDocProcessor
    from rag_chatbot.vector_store import VectorStore
    from rag_chatbot.retriever import Retriever
    from rag_chatbot.chat_engine import ChatEngine, Conversation
    RAG_CHATBOT_AVAILABLE = True
except ImportError as e:
    RAG_CHATBOT_AVAILABLE = False
    print(f"RAG chatbot not available: {e}")

# Import Document Processor
try:
    from Document_processor.processor import DocumentProcessor as AdvancedProcessor
    ADVANCED_PROCESSOR_AVAILABLE = True
except ImportError as e:
    ADVANCED_PROCESSOR_AVAILABLE = False
    print(f"Advanced processor not available: {e}")

# Import Clause Simplification components
try:
    from groq import Groq
    import textstat
    import re
    CLAUSE_SIMPLIFICATION_AVAILABLE = True
except ImportError as e:
    CLAUSE_SIMPLIFICATION_AVAILABLE = False
    print(f"Clause simplification not available: {e}")

# Import NegotiateAI Multi-Agent System
try:
    from negotiate_ai import NegotiateAIOrchestrator
    from negotiate_ai.models import NegotiationPlaybook
    NEGOTIATE_AI_AVAILABLE = True
except ImportError as e:
    NEGOTIATE_AI_AVAILABLE = False
    print(f"NegotiateAI not available: {e}")

# Import Language Translation Module
try:
    from language_translator import LegalTranslator, SUPPORTED_LANGUAGES
    from language_translator.translator import get_language_options
    TRANSLATION_AVAILABLE = True
except ImportError as e:
    TRANSLATION_AVAILABLE = False
    SUPPORTED_LANGUAGES = {"en": {"name": "English", "native_name": "English"}}
    print(f"Translation not available: {e}")

# Import Authentication Module
try:
    from auth import AuthManager
    from auth.auth_manager import login_form, register_form, auth_sidebar
    AUTH_AVAILABLE = True
except ImportError as e:
    AUTH_AVAILABLE = False
    print(f"Auth not available: {e}")

# Import Database Module
try:
    from database import DatabaseManager
    from database.db_manager import get_database
    DATABASE_AVAILABLE = True
except ImportError as e:
    DATABASE_AVAILABLE = False
    print(f"Database not available: {e}")

# Import FAISS Vector Store (cloud-compatible)
try:
    from rag_chatbot.faiss_store import FAISSVectorStore
    FAISS_AVAILABLE = True
except ImportError as e:
    FAISS_AVAILABLE = False
    print(f"FAISS not available: {e}")


# Page configuration
st.set_page_config(
    page_title="ClauseCare - Legal Document Platform",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Custom CSS for all modules
st.markdown("""
<style>
    /* Risk Assessment Styles */
    .risk-critical { 
        background-color: #ffebee; 
        border-left: 4px solid #dc3545; 
        padding: 10px; 
        margin: 5px 0; 
        color: #721c24 !important;
    }
    .risk-high { 
        background-color: #fff3e0; 
        border-left: 4px solid #fd7e14; 
        padding: 10px; 
        margin: 5px 0;
        color: #856404 !important;
    }
    .risk-medium { 
        background-color: #fff8e1; 
        border-left: 4px solid #ffc107; 
        padding: 10px; 
        margin: 5px 0;
        color: #856404 !important;
    }
    .risk-low { 
        background-color: #e8f5e9; 
        border-left: 4px solid #28a745; 
        padding: 10px; 
        margin: 5px 0;
        color: #155724 !important;
    }
    
    .risk-critical strong, .risk-critical span, .risk-critical small { color: #721c24 !important; }
    .risk-high strong, .risk-high span, .risk-high small { color: #856404 !important; }
    .risk-medium strong, .risk-medium span, .risk-medium small { color: #856404 !important; }
    .risk-low strong, .risk-low span, .risk-low small { color: #155724 !important; }
    
    .metric-card {
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 20px;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        color: #333 !important;
    }
    
    .score-gauge {
        font-size: 48px;
        font-weight: bold;
    }
    
    .severity-badge {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 20px;
        color: white;
        font-weight: bold;
    }
    
    .badge-critical { background-color: #dc3545; }
    .badge-high { background-color: #fd7e14; }
    .badge-medium { background-color: #ffc107; color: #333; }
    .badge-low { background-color: #28a745; }
    
    /* Chat Styles - FIXED VISIBILITY */
    .chat-message {
        padding: 1.2rem;
        border-radius: 0.75rem;
        margin-bottom: 1rem;
        color: #1a1a1a !important;
    }
    .chat-message * {
        color: #1a1a1a !important;
    }
    .chat-message.user {
        background-color: #e3f2fd;
        border-left: 4px solid #2196F3;
    }
    .chat-message.user strong {
        color: #1565c0 !important;
    }
    .chat-message.assistant {
        background-color: #f1f8e9;
        border-left: 4px solid #4CAF50;
    }
    .chat-message.assistant strong {
        color: #2e7d32 !important;
    }
    .chat-message h1, .chat-message h2, .chat-message h3, .chat-message h4 {
        color: #333 !important;
        margin-top: 0.5rem;
    }
    .chat-message ul, .chat-message ol, .chat-message li {
        color: #1a1a1a !important;
    }
    .chat-message p {
        color: #1a1a1a !important;
        margin-bottom: 0.5rem;
    }
    
    .source-card {
        background-color: #fff8e1;
        border: 1px solid #ffcc80;
        border-radius: 0.5rem;
        padding: 0.75rem;
        margin: 0.25rem 0;
        font-size: 0.85rem;
        color: #5d4037 !important;
    }
    .source-card * {
        color: #5d4037 !important;
    }
    .source-card strong {
        color: #4e342e !important;
    }
    
    /* Document Info Styles */
    .document-info {
        background-color: #e8eaf6;
        border-radius: 0.5rem;
        padding: 1rem;
        margin-bottom: 1rem;
        color: #283593 !important;
    }
    .document-info * {
        color: #283593 !important;
    }
    
    .legal-term {
        background-color: #fff3e0;
        padding: 0.75rem;
        border-radius: 0.5rem;
        border-left: 3px solid #ff9800;
        margin: 0.5rem 0;
        color: #e65100 !important;
    }
    .legal-term * {
        color: #bf360c !important;
    }
    
    /* Readability Score Styles */
    .readability-card {
        background-color: #f5f5f5;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
        border: 1px solid #ddd;
        color: #333 !important;
    }
    .readability-card * {
        color: #333 !important;
    }
    .readability-score {
        font-size: 32px;
        font-weight: bold;
        text-align: center;
    }
    .readability-easy { color: #28a745 !important; }
    .readability-medium { color: #ffc107 !important; }
    .readability-hard { color: #dc3545 !important; }
    
    /* Simplification Styles */
    .original-clause {
        background-color: #ffebee;
        border-left: 4px solid #ef5350;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
        color: #c62828 !important;
    }
    .original-clause * {
        color: #c62828 !important;
    }
    .simplified-clause {
        background-color: #e8f5e9;
        border-left: 4px solid #66bb6a;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
        color: #2e7d32 !important;
    }
    .simplified-clause * {
        color: #2e7d32 !important;
    }
    
    /* Table Styles */
    .table-container {
        background-color: #f8f9fa;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        overflow-x: auto;
        color: #333 !important;
    }
    
    /* Navigation */
    .nav-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white !important;
        padding: 1rem;
        border-radius: 10px;
        margin-bottom: 1rem;
    }
    .nav-header * {
        color: white !important;
    }
    
    /* Analysis Card */
    .analysis-card {
        background-color: #fafafa;
        border: 1px solid #e0e0e0;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
        color: #333 !important;
    }
    .analysis-card * {
        color: #333 !important;
    }
    
    /* ================================
       RESPONSIVE DESIGN - Mobile First
       ================================ */
    
    /* Base responsive adjustments */
    .stApp {
        max-width: 100%;
        overflow-x: hidden;
    }
    
    /* Mobile devices (portrait) */
    @media screen and (max-width: 640px) {
        /* Reduce padding for mobile */
        .block-container {
            padding-left: 1rem !important;
            padding-right: 1rem !important;
            padding-top: 1rem !important;
        }
        
        /* Stack columns vertically */
        [data-testid="column"] {
            width: 100% !important;
            flex: 1 1 100% !important;
            min-width: 100% !important;
        }
        
        /* Adjust font sizes */
        .nav-header h1 {
            font-size: 1.5rem !important;
        }
        
        .score-gauge {
            font-size: 32px !important;
        }
        
        .metric-card {
            padding: 12px !important;
        }
        
        /* Risk cards mobile */
        .risk-critical, .risk-high, .risk-medium, .risk-low {
            padding: 8px !important;
            margin: 4px 0 !important;
            font-size: 14px !important;
        }
        
        /* Chat messages mobile */
        .chat-message {
            padding: 0.75rem !important;
            margin-bottom: 0.5rem !important;
        }
        
        /* Tables scroll horizontally */
        .table-container {
            overflow-x: auto !important;
            -webkit-overflow-scrolling: touch;
        }
        
        /* Sidebar adjustments */
        [data-testid="stSidebar"] {
            min-width: 200px !important;
        }
        
        /* Feature cards mobile */
        .feature-card {
            padding: 15px !important;
            margin: 8px 0 !important;
        }
        
        /* Expander content */
        .streamlit-expanderContent {
            padding: 8px !important;
        }
        
        /* Form inputs */
        .stTextInput > div > div > input {
            font-size: 16px !important; /* Prevents zoom on iOS */
        }
        
        /* Buttons touch-friendly */
        .stButton > button {
            min-height: 44px !important;
            padding: 8px 16px !important;
        }
        
        /* Download buttons */
        .stDownloadButton > button {
            width: 100% !important;
        }
    }
    
    /* Tablets (portrait) */
    @media screen and (min-width: 641px) and (max-width: 1024px) {
        .block-container {
            padding-left: 2rem !important;
            padding-right: 2rem !important;
        }
        
        /* 2-column layout on tablets */
        [data-testid="column"] {
            min-width: 45% !important;
        }
        
        .score-gauge {
            font-size: 40px !important;
        }
    }
    
    /* Large screens */
    @media screen and (min-width: 1025px) {
        .block-container {
            max-width: 1200px !important;
            margin: 0 auto !important;
        }
    }
    
    /* Print styles */
    @media print {
        [data-testid="stSidebar"] {
            display: none !important;
        }
        
        .stButton, .stDownloadButton {
            display: none !important;
        }
        
        .block-container {
            padding: 0 !important;
            max-width: 100% !important;
        }
    }
    
    /* Touch device optimizations */
    @media (hover: none) and (pointer: coarse) {
        /* Larger touch targets */
        .stButton > button,
        .stDownloadButton > button,
        [data-testid="stSidebar"] button {
            min-height: 48px !important;
            min-width: 48px !important;
        }
        
        /* More spacing between interactive elements */
        .stRadio > div,
        .stCheckbox > div {
            gap: 12px !important;
        }
    }
    
    /* Feature card grid */
    .feature-card {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        border-radius: 12px;
        padding: 20px;
        margin: 10px 0;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
        color: #333 !important;
    }
    .feature-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.15);
    }
    .feature-card * {
        color: #333 !important;
    }
    .feature-card h3 {
        color: #1a237e !important;
    }
    
    /* Auth form styles */
    .auth-container {
        max-width: 400px;
        margin: 0 auto;
        padding: 20px;
        background: #fff;
        border-radius: 12px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    /* User badge in sidebar */
    .user-badge {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white !important;
        padding: 10px 15px;
        border-radius: 8px;
        margin-bottom: 10px;
    }
    .user-badge * {
        color: white !important;
    }
</style>
""", unsafe_allow_html=True)


# ============================================
# UTILITY FUNCTIONS
# ============================================

def get_severity_color(severity: str) -> str:
    """Get color for severity level"""
    colors = {
        "CRITICAL": "#dc3545",
        "HIGH": "#fd7e14",
        "MEDIUM": "#ffc107",
        "LOW": "#28a745",
    }
    return colors.get(severity, "#6c757d")


def get_severity_emoji(severity: str) -> str:
    """Get emoji for severity level"""
    emojis = {
        "CRITICAL": "🚨",
        "HIGH": "🔶",
        "MEDIUM": "⚠️",
        "LOW": "✅",
    }
    return emojis.get(severity, "❓")


def extract_text_from_file(uploaded_file) -> tuple[str, dict]:
    """
    Extract text from uploaded file using best available method.
    Returns (text, metadata)
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)  # Reset for other uses
    
    text = ""
    metadata = {"filename": uploaded_file.name, "size": len(file_bytes)}
    
    # Try advanced processor first if available
    if ADVANCED_PROCESSOR_AVAILABLE:
        try:
            processor = AdvancedProcessor()
            uploaded_file.seek(0)
            result = processor.process_document(uploaded_file)
            
            if result.get("status") == "success":
                text = result.get("text", "")
                metadata.update(result.get("metadata", {}))
                metadata["tables"] = result.get("tables", [])
                metadata["analysis"] = result.get("analysis", {})
                return text, metadata
        except Exception as e:
            print(f"Advanced processor failed: {e}")
    
    # Fallback to basic extraction
    if filename.endswith('.txt'):
        text = file_bytes.decode('utf-8')
        metadata["type"] = "TXT"
        
    elif filename.endswith('.pdf'):
        try:
            import fitz
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)
            text = "\n".join(text_parts)
            metadata["pages"] = len(pdf_doc)
            metadata["type"] = "PDF"
            pdf_doc.close()
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {e}")
            
    elif filename.endswith('.docx'):
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            metadata["paragraphs"] = len(doc.paragraphs)
            metadata["type"] = "DOCX"
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {e}")
    
    return text, metadata


def init_session_state():
    """Initialize all session state variables"""
    defaults = {
        # Authentication state
        'authenticated': False,
        'current_user': None,
        'auth_message': None,
        
        # Document state
        'document_text': None,
        'document_name': None,
        'document_metadata': {},
        'document_loaded': False,
        
        # Risk Assessment state
        'analysis_result': None,
        
        # RAG Chatbot state
        'vector_store': None,
        'retriever': None,
        'chat_engine': None,
        'conversation': None,
        'chat_messages': [],
        'suggested_questions': [],
        'rag_initialized': False,
        
        # Clause Simplification state
        'clauses': [],
        'simplified_clauses': {},
        
        # NegotiateAI state
        'negotiation_playbook': None,
        'negotiation_context': {},
        'negotiation_in_progress': False,
        'negotiation_progress': {},
        
        # Language Translation state
        'selected_language': 'en',
        'translated_content': {},
        'translation_history': [],
        
        # Navigation
        'current_page': 'home',
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def initialize_rag_system():
    """Initialize RAG system components using FAISS (cloud-compatible)"""
    if not RAG_CHATBOT_AVAILABLE:
        return False
    
    if not GROQ_API_KEY:
        return False
    
    try:
        # Use FAISS for cloud deployment (no persistence needed)
        if st.session_state.vector_store is None:
            if FAISS_AVAILABLE:
                st.session_state.vector_store = FAISSVectorStore()
            else:
                # Fallback to ChromaDB for local development
                st.session_state.vector_store = VectorStore(
                    collection_name="clausecare_docs",
                    persist_directory="./chroma_db"
                )
        
        if st.session_state.retriever is None:
            st.session_state.retriever = Retriever(
                vector_store=st.session_state.vector_store,
                max_context_tokens=4000,
                top_k=5,
            )
        
        if st.session_state.chat_engine is None:
            st.session_state.chat_engine = ChatEngine(
                retriever=st.session_state.retriever,
                api_key=GROQ_API_KEY,
            )
        
        st.session_state.rag_initialized = True
        return True
        
    except Exception as e:
        print(f"RAG initialization failed: {e}")
        return False


# ============================================
# RISK ASSESSMENT COMPONENTS
# ============================================

def render_risk_gauge(score: int, label: str = "Overall Risk"):
    """Render a risk score gauge"""
    if score <= 30:
        color = "#28a745"
        level = "LOW"
    elif score <= 60:
        color = "#ffc107"
        level = "MEDIUM"
    elif score <= 85:
        color = "#fd7e14"
        level = "HIGH"
    else:
        color = "#dc3545"
        level = "CRITICAL"
    
    st.markdown(f"""
    <div class="metric-card">
        <div style="color: {color};" class="score-gauge">{score}</div>
        <div style="font-size: 14px; color: #666;">{label}</div>
        <div style="margin-top: 10px;">
            <span class="severity-badge badge-{level.lower()}">{level}</span>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_distribution_metrics(distribution):
    """Render distribution metrics in columns"""
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("🚨 Critical", distribution.critical)
    with col2:
        st.metric("🔶 High", distribution.high)
    with col3:
        st.metric("⚠️ Medium", distribution.medium)
    with col4:
        st.metric("✅ Low", distribution.low)


def render_top_risks(top_risks):
    """Render top risks as expandable sections"""
    for risk in top_risks[:5]:
        severity = "CRITICAL" if risk.score >= 85 else "HIGH" if risk.score >= 60 else "MEDIUM"
        emoji = get_severity_emoji(severity)
        
        with st.expander(f"{emoji} {risk.clause_reference} (Score: {risk.score}/100)"):
            st.markdown(f"**Issue:** {risk.issue}")
            st.markdown(f"**Recommended Action:** {risk.action}")
            st.progress(risk.score / 100)


def render_clause_risks(clause_risks):
    """Render all clause risks in a tabular format"""
    for risk in clause_risks:
        severity = risk.severity.value
        css_class = f"risk-{severity.lower()}"
        emoji = get_severity_emoji(severity)
        
        clause_type_display = risk.clause_type.replace("_", " ").title()
        if risk.clause_type == "unanalyzed":
            clause_type_display = risk.category.value.replace("_", " ").title() + " (Rule-Based)"
        
        ai_status = "🤖" if risk.ai_analyzed else "📋"
        
        st.markdown(f"""
        <div class="{css_class}">
            <strong>{emoji} {risk.clause_id}</strong> - {clause_type_display} {ai_status}
            <span style="float: right;">Score: {risk.score}/100</span>
            <br><small>{risk.primary_risk if risk.ai_analyzed else f"Potential risk detected in {risk.category.value.replace('_', ' ')} clause"}</small>
        </div>
        """, unsafe_allow_html=True)
        
        with st.expander("View Details"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Category:**")
                st.write(risk.category.value.replace("_", " ").title())
                
                st.markdown("**Confidence:**")
                st.progress(risk.confidence / 100)
                st.write(f"{risk.confidence}%")
                
                if risk.red_flags:
                    st.markdown("**🚩 Red Flags:**")
                    for flag in risk.red_flags[:5]:
                        st.write(f"• {flag}")
            
            with col2:
                st.markdown("**Detailed Explanation:**")
                st.write(risk.detailed_explanation or risk.primary_risk)
                
                st.markdown("**Recommendation:**")
                st.info(risk.recommendation)
                
                if risk.alternative_language:
                    st.markdown("**Suggested Alternative:**")
                    st.code(risk.alternative_language)


def render_action_plan(document_risk):
    """Render the action plan"""
    st.subheader("📋 Action Plan")
    for i, action in enumerate(document_risk.action_plan, 1):
        st.markdown(f"{action}")


def run_risk_assessment():
    """Run the Risk Assessment page"""
    st.title("🔍 Risk Assessment")
    st.markdown("Analyze your legal document for potential risks and get actionable recommendations.")
    
    if not RISK_ASSESSMENT_AVAILABLE:
        st.error("Risk Assessment module not available. Please check installation.")
        return
    
    # Document context settings
    col1, col2, col3 = st.columns(3)
    
    with col1:
        doc_type = st.selectbox(
            "Document Type",
            ["service_agreement", "nda", "employment_contract", 
             "lease_agreement", "partnership_agreement", "other"]
        )
    
    with col2:
        user_role = st.selectbox(
            "Your Role",
            ["customer", "service_provider", "employee", "employer",
             "tenant", "landlord", "party_reviewing"]
        )
    
    with col3:
        industry = st.selectbox(
            "Industry",
            ["technology", "healthcare", "finance", "retail",
             "manufacturing", "consulting", "general"]
        )
    
    use_ai = st.checkbox("Enable AI Analysis", value=True, help="Use Groq Llama 3.3 70B for deep analysis")
    
    # Check if document is loaded
    if not st.session_state.document_loaded:
        st.warning("⚠️ Please upload a document first from the sidebar.")
        return
    
    st.info(f"📄 Analyzing: **{st.session_state.document_name}** ({len(st.session_state.document_text):,} characters)")
    
    # Analysis button
    if st.button("🔍 Analyze Document", type="primary", use_container_width=True):
        context = AnalysisContext(
            document_type=doc_type,
            user_role=user_role,
            industry=industry,
        )
        
        engine = RiskAssessmentEngine(
            api_key=GROQ_API_KEY if GROQ_API_KEY else None,
            use_ai=use_ai and bool(GROQ_API_KEY),
        )
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        def update_progress(current, total, message):
            progress_bar.progress(current / total)
            status_text.text(message)
        
        callback = ProgressCallback(
            on_start=lambda msg: status_text.text(msg),
            on_progress=update_progress,
            on_complete=lambda msg: status_text.success(msg),
        )
        
        with st.spinner("Analyzing document..."):
            try:
                result = engine.analyze_document(
                    text=st.session_state.document_text,
                    filename=st.session_state.document_name,
                    context=context,
                    progress_callback=callback,
                )
                st.session_state['analysis_result'] = result
                st.session_state['risk_engine'] = engine
            except Exception as e:
                st.error(f"Analysis failed: {str(e)}")
                return
    
    # Display results
    if st.session_state.get('analysis_result'):
        result = st.session_state['analysis_result']
        
        st.markdown("---")
        st.header("📊 Risk Assessment Results")
        
        # Summary metrics row
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col1:
            render_risk_gauge(result.risk_summary.overall_score)
        
        with col2:
            st.subheader("Executive Summary")
            st.markdown(result.executive_summary)
            
            favorability_colors = {
                "heavily_favors_other_party": "🔴",
                "slightly_unfavorable": "🟠",
                "balanced": "🟡",
                "favorable": "🟢",
            }
            fav_emoji = favorability_colors.get(result.overall_favorability, "⚪")
            st.markdown(f"**Contract Favorability:** {fav_emoji} {result.overall_favorability.replace('_', ' ').title()}")
        
        with col3:
            st.subheader("Quick Stats")
            st.metric("Clauses Analyzed", result.metadata.clauses_analyzed)
            st.metric("Processing Time", f"{result.metadata.processing_time_seconds:.1f}s")
        
        # Distribution metrics
        st.subheader("Risk Distribution")
        render_distribution_metrics(result.risk_summary.distribution)
        
        # Tabs for different views
        tab_overview, tab_details, tab_actions, tab_export = st.tabs([
            "📈 Overview", "📋 Clause Details", "🎯 Action Items", "📥 Export"
        ])
        
        with tab_overview:
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("🔝 Top Risks")
                render_top_risks(result.top_risks)
            
            with col2:
                st.subheader("📊 Risk by Category")
                if result.category_summaries:
                    category_data = {
                        cat.value.replace("_", " ").title(): summary.average_score
                        for cat, summary in result.category_summaries.items()
                    }
                    st.bar_chart(category_data)
                else:
                    st.info("No category data available")
            
            if result.deal_breakers:
                st.subheader("🚫 Potential Deal Breakers")
                for db in result.deal_breakers:
                    st.error(f"⚠️ {db}")
        
        with tab_details:
            st.subheader("All Clause Assessments")
            
            severity_filter = st.multiselect(
                "Filter by Severity",
                ["CRITICAL", "HIGH", "MEDIUM", "LOW"],
                default=["CRITICAL", "HIGH", "MEDIUM"]
            )
            
            filtered_risks = [
                r for r in result.clause_risks
                if r.severity.value in severity_filter
            ]
            
            if filtered_risks:
                render_clause_risks(filtered_risks)
            else:
                st.info("No clauses match the selected filter.")
        
        with tab_actions:
            if result.must_address_immediately:
                st.subheader("🚨 Must Address Immediately")
                for action in result.must_address_immediately:
                    st.error(f"""
                    **{action.clause_reference}**
                    
                    Issue: {action.issue}
                    
                    Action: {action.action}
                    """)
            
            if result.should_negotiate:
                st.subheader("⚠️ Should Negotiate")
                for clause in result.should_negotiate:
                    st.warning(clause)
            
            render_action_plan(result)
            
            if result.acceptable_as_is:
                st.subheader("✅ Acceptable Terms")
                st.success(", ".join(result.acceptable_as_is[:10]))
        
        with tab_export:
            st.subheader("Export Report")
            
            col1, col2 = st.columns(2)
            
            engine = st.session_state.get('risk_engine')
            if engine:
                with col1:
                    markdown_report = engine.aggregator.to_markdown_report(result)
                    st.download_button(
                        label="📄 Download Markdown Report",
                        data=markdown_report,
                        file_name="risk_assessment_report.md",
                        mime="text/markdown",
                        use_container_width=True,
                    )
                
                with col2:
                    json_report = json.dumps(engine.aggregator.to_json(result), indent=2)
                    st.download_button(
                        label="📊 Download JSON Data",
                        data=json_report,
                        file_name="risk_assessment_data.json",
                        mime="application/json",
                        use_container_width=True,
                    )
                
                st.subheader("Report Preview")
                st.markdown(markdown_report[:2000] + "..." if len(markdown_report) > 2000 else markdown_report)


# ============================================
# RAG CHATBOT COMPONENTS
# ============================================

def render_chat_message(role: str, content: str):
    """Render a chat message"""
    css_class = "user" if role == "user" else "assistant"
    icon = "🧑" if role == "user" else "🤖"
    
    st.markdown(f"""
    <div class="chat-message {css_class}">
        <strong>{icon} {'You' if role == 'user' else 'Assistant'}</strong><br>
        {content}
    </div>
    """, unsafe_allow_html=True)


def render_sources(sources: list):
    """Render source citations"""
    if not sources:
        return
    
    with st.expander("📄 View Sources", expanded=False):
        for source in sources:
            section = source.get('section', 'Unknown')
            page = source.get('page_number', '?')
            snippet = source.get('snippet', '')[:150] + "..."
            
            st.markdown(f"""
            <div class="source-card">
                <strong>[Source {source.get('index', '?')}]</strong> {section} (Page {page})<br>
                <em>"{snippet}"</em>
            </div>
            """, unsafe_allow_html=True)


def run_document_chat():
    """Run the Document Chat page"""
    st.title("💬 Document Chat")
    st.markdown("Ask questions about your legal document and get answers with citations.")
    
    if not RAG_CHATBOT_AVAILABLE:
        st.error("RAG Chatbot module not available. Please check installation.")
        return
    
    if not GROQ_API_KEY:
        st.error("⚠️ GROQ_API_KEY not found in environment. Please set it in your .env file.")
        return
    
    if not st.session_state.document_loaded:
        st.warning("⚠️ Please upload a document first from the sidebar.")
        
        st.markdown("### Once you upload a document, you can ask questions like:")
        example_questions = [
            "What are my main obligations under this contract?",
            "What happens if I want to terminate early?",
            "What are the payment terms and deadlines?",
            "Who is liable if something goes wrong?",
            "What information do I need to keep confidential?",
        ]
        for q in example_questions:
            st.markdown(f"- _{q}_")
        return
    
    # Initialize RAG system
    if not st.session_state.rag_initialized:
        with st.spinner("Initializing chat engine..."):
            if not initialize_rag_system():
                st.error("Failed to initialize chat engine.")
                return
    
    # Check if document is indexed
    if st.session_state.conversation is None:
        with st.spinner("Indexing document for chat..."):
            try:
                # Process document for RAG
                rag_processor = RAGDocProcessor(chunk_size=500, chunk_overlap=100)
                doc = rag_processor.process_text(
                    st.session_state.document_text,
                    st.session_state.document_name
                )
                
                # Clear old data and add new
                st.session_state.vector_store.clear()
                chunks_added = st.session_state.vector_store.add_document(doc)
                
                # Create conversation
                st.session_state.conversation = st.session_state.chat_engine.create_conversation(
                    document_id=doc.doc_id,
                    document_name=st.session_state.document_name,
                )
                
                st.session_state.suggested_questions = [
                    "What is this document about?",
                    "What are my main obligations?",
                    "What are the termination conditions?",
                    "What are the payment terms?",
                    "Are there any liability limitations?",
                ]
                
                st.success(f"✅ Indexed {chunks_added} chunks from document")
                
            except Exception as e:
                st.error(f"Failed to index document: {str(e)}")
                return
    
    # Document info
    st.markdown(f"""
    <div class="document-info">
        📄 <strong>Chatting about:</strong> {st.session_state.document_name}
    </div>
    """, unsafe_allow_html=True)
    
    # Chat messages
    for msg in st.session_state.chat_messages:
        render_chat_message(msg["role"], msg["content"])
        if msg.get("sources"):
            render_sources(msg["sources"])
    
    # Suggested questions
    if not st.session_state.chat_messages and st.session_state.suggested_questions:
        st.markdown("**💡 Suggested questions:**")
        cols = st.columns(min(len(st.session_state.suggested_questions), 3))
        
        for i, question in enumerate(st.session_state.suggested_questions[:6]):
            col_idx = i % 3
            with cols[col_idx]:
                if st.button(question, key=f"suggested_{i}", use_container_width=True):
                    st.session_state.pending_question = question
                    st.rerun()
    
    # Chat input
    st.markdown("---")
    
    # Check for pending question
    pending = st.session_state.get('pending_question')
    prompt = st.chat_input("Ask a question about your document...")
    
    if pending:
        prompt = pending
        del st.session_state.pending_question
    
    if prompt:
        # Add user message
        st.session_state.chat_messages.append({"role": "user", "content": prompt})
        render_chat_message("user", prompt)
        
        # Generate response
        with st.spinner("Thinking..."):
            response_placeholder = st.empty()
            full_response = ""
            sources = []
            
            try:
                for chunk in st.session_state.chat_engine.chat(
                    query=prompt,
                    conversation=st.session_state.conversation,
                    stream=True,
                ):
                    full_response += chunk
                    response_placeholder.markdown(full_response + "▌")
                
                response_placeholder.markdown(full_response)
                
                if st.session_state.conversation.messages:
                    last_msg = st.session_state.conversation.messages[-1]
                    sources = last_msg.sources
                
                st.session_state.chat_messages.append({
                    "role": "assistant",
                    "content": full_response,
                    "sources": sources,
                })
                
                st.session_state.suggested_questions = \
                    st.session_state.chat_engine.get_suggested_questions(
                        st.session_state.conversation
                    )
                
            except Exception as e:
                st.error(f"Error generating response: {str(e)}")
        
        st.rerun()
    
    # Legal term lookup
    st.markdown("---")
    with st.expander("📚 Look up a legal term"):
        term = st.text_input("Enter a legal term", placeholder="e.g., indemnification")
        if term:
            explanation = st.session_state.chat_engine.explain_term(term)
            st.markdown(f"""
            <div class="legal-term">
                {explanation}
            </div>
            """, unsafe_allow_html=True)


# ============================================
# DOCUMENT PROCESSOR PAGE
# ============================================

def run_document_processor():
    """Run the Document Processor page"""
    st.title("📄 Document Processor")
    st.markdown("Advanced document analysis with OCR, table extraction, and text analysis.")
    
    if not st.session_state.document_loaded:
        st.warning("⚠️ Please upload a document first from the sidebar.")
        return
    
    metadata = st.session_state.document_metadata
    
    # Document info
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("📄 Document", st.session_state.document_name)
    with col2:
        st.metric("📝 Characters", f"{len(st.session_state.document_text):,}")
    with col3:
        doc_type = metadata.get("type", "Unknown")
        st.metric("📋 Type", doc_type)
    
    # Metadata
    st.subheader("📊 Document Metadata")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Basic Information:**")
        for key, value in metadata.items():
            if key not in ['tables', 'analysis', 'text']:
                st.write(f"- **{key.title()}:** {value}")
    
    with col2:
        if "analysis" in metadata and metadata["analysis"]:
            st.markdown("**Text Analysis:**")
            analysis = metadata["analysis"]
            st.write(f"- **Language:** {analysis.get('language', 'Unknown')}")
            st.write(f"- **Readability Score:** {analysis.get('readability_score', 'N/A')}")
            st.write(f"- **Difficulty:** {analysis.get('difficulty', 'N/A')}")
    
    # Tables
    if "tables" in metadata and metadata["tables"]:
        st.subheader("📊 Extracted Tables")
        
        for table in metadata["tables"]:
            with st.expander(f"Table {table.get('table_number', '?')} (Page {table.get('page', '?')})"):
                st.write(f"**Rows:** {table.get('rows', 0)} | **Columns:** {table.get('columns', 0)}")
                
                if table.get('data'):
                    try:
                        import pandas as pd
                        df = pd.DataFrame(table['data'][1:], columns=table['data'][0] if table['data'] else [])
                        st.dataframe(df, use_container_width=True)
                    except:
                        st.code(table.get('text_representation', 'No data'))
    
    # Document text preview
    st.subheader("📝 Document Text")
    
    with st.expander("View Full Text", expanded=False):
        st.text_area(
            "Document Content",
            value=st.session_state.document_text,
            height=400,
            disabled=True
        )
    
    # Export options
    st.subheader("📥 Export Options")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.download_button(
            "📄 Download as TXT",
            data=st.session_state.document_text,
            file_name=f"{st.session_state.document_name.rsplit('.', 1)[0]}.txt",
            mime="text/plain",
            use_container_width=True,
        )
    
    with col2:
        metadata_json = json.dumps(metadata, indent=2, default=str)
        st.download_button(
            "📊 Download Metadata JSON",
            data=metadata_json,
            file_name="document_metadata.json",
            mime="application/json",
            use_container_width=True,
        )


# ============================================
# CLAUSE SIMPLIFICATION COMPONENTS
# ============================================

def split_clauses(text: str) -> list[str]:
    """Split legal text into numbered clauses"""
    # Detect numbered patterns like 1., 1.1, (a), Article 1, Section 5
    pattern = r'(?:\n\s*(\d+\.|[a-z]\)|\d+\.\d+|ARTICLE [IVX\d]+|SECTION \d+)\s+)'
    parts = re.split(pattern, text) if CLAUSE_SIMPLIFICATION_AVAILABLE else [text]
    clauses = []
    
    if len(parts) > 1:
        for i in range(1, len(parts), 2):
            marker = parts[i]
            content = parts[i+1] if i+1 < len(parts) else ""
            full_clause = f"{marker} {content}".strip()
            if len(full_clause) > 20:
                clauses.append(full_clause)
    else:
        # Fallback: split by paragraphs
        clauses = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 20]
    
    return clauses[:50]  # Limit to 50 clauses


def get_readability_score(text: str) -> tuple[float, str, str]:
    """
    Get Flesch Reading Ease score and interpretation.
    Returns (score, level, css_class)
    """
    if not CLAUSE_SIMPLIFICATION_AVAILABLE or not text.strip():
        return 0, "Unknown", "readability-medium"
    
    try:
        score = textstat.flesch_reading_ease(text)
        
        if score >= 80:
            return score, "Very Easy (5th Grade)", "readability-easy"
        elif score >= 60:
            return score, "Easy (8th Grade)", "readability-easy"
        elif score >= 40:
            return score, "Standard (High School)", "readability-medium"
        elif score >= 20:
            return score, "Difficult (College)", "readability-hard"
        else:
            return score, "Very Difficult (Legal/Academic)", "readability-hard"
    except:
        return 50, "Standard", "readability-medium"


def simplify_clause_with_groq(clause_text: str) -> str:
    """Use Groq to simplify a legal clause"""
    if not GROQ_API_KEY:
        return "Error: Groq API key not configured"
    
    try:
        client = Groq(api_key=GROQ_API_KEY)
        
        prompt = f"""You are a legal expert explaining things to a 10-year-old.
Rewrite the following legal clause in simple, everyday language.

Rules:
1. Use very simple words (avoid legal jargon)
2. Keep sentences short (max 15 words each)
3. Start with "In simple terms:"
4. Explain what this means for a regular person
5. Keep your response under 150 words

Legal Clause:
{clause_text}

Simplified Version:"""
        
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.3
        )
        
        return response.choices[0].message.content or "Unable to simplify"
        
    except Exception as e:
        return f"Error: {str(e)}"


def run_clause_simplification():
    """Run the Clause Simplification page"""
    st.title("✨ Clause Simplifier")
    st.markdown("Transform complex legal jargon into plain, understandable English.")
    
    if not CLAUSE_SIMPLIFICATION_AVAILABLE:
        st.error("Clause Simplification module not available. Install: `pip install groq textstat`")
        return
    
    if not GROQ_API_KEY:
        st.error("⚠️ GROQ_API_KEY not found. Please set it in your .env file.")
        return
    
    if not st.session_state.document_loaded:
        st.warning("⚠️ Please upload a document first from the sidebar.")
        return
    
    # Extract clauses if not done
    if not st.session_state.clauses:
        with st.spinner("Extracting clauses from document..."):
            st.session_state.clauses = split_clauses(st.session_state.document_text)
    
    clauses = st.session_state.clauses
    
    if not clauses:
        st.warning("Could not extract clauses from document. Try a different document.")
        return
    
    # Document readability overview
    st.subheader("📊 Document Readability Analysis")
    
    doc_score, doc_level, doc_css = get_readability_score(st.session_state.document_text)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown(f"""
        <div class="readability-card">
            <div class="readability-score {doc_css}">{doc_score:.1f}</div>
            <div style="text-align: center; margin-top: 10px;">
                <strong>Flesch Reading Ease</strong><br>
                <span>{doc_level}</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.metric("📄 Total Clauses", len(clauses))
    
    with col3:
        simplified_count = len(st.session_state.get('simplified_clauses', {}))
        st.metric("✅ Simplified", simplified_count)
    
    st.markdown("---")
    
    # Clause selector
    st.subheader("🔍 Select a Clause to Simplify")
    
    selected_idx = st.selectbox(
        "Choose a clause:",
        range(len(clauses)),
        format_func=lambda x: f"Clause {x+1}: {clauses[x][:80]}..."
    )
    
    selected_clause = clauses[selected_idx]
    
    # Two column layout
    col_orig, col_simp = st.columns(2)
    
    with col_orig:
        st.markdown("### 📜 Original Legal Text")
        
        orig_score, orig_level, orig_css = get_readability_score(selected_clause)
        
        st.markdown(f"""
        <div class="original-clause">
            {selected_clause[:1500]}{'...' if len(selected_clause) > 1500 else ''}
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <div class="readability-card">
            <strong>Readability Score:</strong> 
            <span class="{orig_css}">{orig_score:.1f}</span> - {orig_level}
        </div>
        """, unsafe_allow_html=True)
    
    with col_simp:
        st.markdown("### ✨ AI Simplified Version")
        
        # Check if already simplified
        cached_key = f"clause_{selected_idx}"
        
        if st.button("🚀 Simplify This Clause", type="primary", use_container_width=True):
            with st.spinner("AI is rewriting in simple terms..."):
                start_time = time.time()
                simplified = simplify_clause_with_groq(selected_clause)
                end_time = time.time()
                
                st.session_state.simplified_clauses[cached_key] = {
                    'text': simplified,
                    'time': end_time - start_time
                }
        
        if cached_key in st.session_state.get('simplified_clauses', {}):
            cached = st.session_state.simplified_clauses[cached_key]
            simplified_text = cached['text']
            gen_time = cached['time']
            
            st.markdown(f"""
            <div class="simplified-clause">
                {simplified_text}
            </div>
            """, unsafe_allow_html=True)
            
            # New readability score
            new_score, new_level, new_css = get_readability_score(simplified_text)
            
            st.markdown(f"""
            <div class="readability-card">
                <strong>New Readability Score:</strong> 
                <span class="{new_css}">{new_score:.1f}</span> - {new_level}
            </div>
            """, unsafe_allow_html=True)
            
            # Improvement metrics
            improvement = new_score - orig_score
            if improvement > 0:
                st.success(f"📈 Readability improved by **{improvement:.1f} points!**")
            
            st.caption(f"⏱️ Generated in {gen_time:.2f}s using Llama 3.3 70B")
        else:
            st.info("👆 Click the button above to simplify this clause")
    
    # Batch simplification option
    st.markdown("---")
    with st.expander("⚡ Batch Simplify All Clauses"):
        st.warning("This will use your API quota to simplify all clauses. May take several minutes.")
        
        if st.button("🔄 Simplify All Clauses"):
            progress = st.progress(0)
            status = st.empty()
            
            for i, clause in enumerate(clauses[:20]):  # Limit to 20
                cache_key = f"clause_{i}"
                if cache_key not in st.session_state.simplified_clauses:
                    status.text(f"Simplifying clause {i+1}/{min(len(clauses), 20)}...")
                    simplified = simplify_clause_with_groq(clause)
                    st.session_state.simplified_clauses[cache_key] = {
                        'text': simplified,
                        'time': 0
                    }
                progress.progress((i + 1) / min(len(clauses), 20))
                time.sleep(0.5)  # Rate limiting
            
            status.success("✅ Batch simplification complete!")
            st.rerun()


# ============================================
# NEGOTIATE AI - MULTI-AGENT SYSTEM
# ============================================

def run_negotiate_ai():
    """Run the NegotiateAI Multi-Agent Negotiation Intelligence page"""
    st.markdown("""
    <div class="nav-header">
        <h1>🤝 NegotiateAI</h1>
        <p>Multi-Agent Contract Negotiation Intelligence</p>
    </div>
    """, unsafe_allow_html=True)
    
    if not NEGOTIATE_AI_AVAILABLE:
        st.error("❌ NegotiateAI module not available. Please check installation.")
        return
    
    if not GROQ_API_KEY:
        st.error("❌ Please set GROQ_API_KEY in your .env file")
        return
    
    if not st.session_state.document_loaded:
        st.info("📄 Please upload a legal document using the sidebar to begin negotiation analysis.")
        
        # Show what NegotiateAI does
        st.markdown("""
        ### 🎯 What NegotiateAI Does
        
        NegotiateAI deploys a team of 6 specialized AI agents to analyze your contract:
        
        | Agent | Role | What They Do |
        |-------|------|--------------|
        | 📋 Document Analyzer | Structure Expert | Extracts all clauses, parties, obligations |
        | ⚠️ Risk Assessor | Risk Specialist | Identifies and scores every risk (0-100) |
        | 🎯 Negotiation Strategist | Master Negotiator | Develops tactical playbook with talking points |
        | ⚖️ Legal Advisor | Legal Counsel | Flags compliance issues, cites precedents |
        | 📊 Market Researcher | Market Intel | Benchmarks terms against industry standards |
        | 🔧 Contract Optimizer | Chief Synthesizer | Creates prioritized action plan |
        
        **Result:** A complete negotiation playbook that gives you the intelligence of a Fortune 500 legal team.
        """)
        return
    
    # Document is loaded - show analysis options
    st.success(f"📄 Document loaded: **{st.session_state.document_name}**")
    
    # Context configuration
    with st.expander("⚙️ Configure Analysis Context", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            your_role = st.selectbox(
                "Your Role in This Contract",
                ["Service Provider", "Client/Customer", "Vendor", "Contractor", 
                 "Employer", "Employee", "Landlord", "Tenant", "Partner", "Other"],
                help="Select your role to get personalized negotiation strategies"
            )
            
            industry = st.selectbox(
                "Industry",
                ["Technology/SaaS", "Healthcare", "Financial Services", "Manufacturing",
                 "Retail/E-commerce", "Real Estate", "Professional Services", 
                 "Media/Entertainment", "Energy", "Other"],
                help="Industry context for market benchmarking"
            )
        
        with col2:
            jurisdiction = st.selectbox(
                "Primary Jurisdiction",
                ["United States", "California", "New York", "Texas", "Delaware",
                 "United Kingdom", "European Union", "Canada", "Australia", "Other"],
                help="Legal jurisdiction for compliance analysis"
            )
            
            contract_value = st.text_input(
                "Estimated Contract Value",
                placeholder="e.g., $100,000/year",
                help="Helps with market benchmarking and risk assessment"
            )
        
        importance = st.radio(
            "Deal Importance",
            ["Standard business deal", "Strategic partnership", "Critical/Must-win", "Exploratory"],
            horizontal=True
        )
    
    # Store context
    st.session_state.negotiation_context = {
        "your_role": your_role,
        "industry": industry.split("/")[0],  # Get primary category
        "jurisdiction": jurisdiction,
        "contract_value": contract_value or "Not specified",
        "importance": importance
    }
    
    # Analysis buttons
    st.markdown("---")
    
    col1, col2, col3 = st.columns([2, 2, 1])
    
    with col1:
        run_full = st.button(
            "🚀 Run Full 6-Agent Analysis",
            type="primary",
            use_container_width=True,
            help="Comprehensive analysis by all 6 agents (~2-3 minutes)"
        )
    
    with col2:
        if st.session_state.negotiation_playbook:
            st.download_button(
                "📥 Download Playbook",
                data=st.session_state.negotiation_playbook.executive_summary,
                file_name=f"negotiation_playbook_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                mime="text/markdown",
                use_container_width=True
            )
    
    with col3:
        if st.session_state.negotiation_playbook:
            if st.button("🗑️ Clear", use_container_width=True):
                st.session_state.negotiation_playbook = None
                st.rerun()
    
    # Run analysis
    if run_full:
        run_full_negotiation_analysis()
    
    # Display results
    if st.session_state.negotiation_playbook:
        display_negotiation_playbook(st.session_state.negotiation_playbook)


def run_full_negotiation_analysis():
    """Run the full 6-agent analysis pipeline"""
    orchestrator = NegotiateAIOrchestrator(api_key=GROQ_API_KEY)
    
    # Progress display
    progress_container = st.container()
    
    with progress_container:
        st.markdown("### 🔄 Analysis in Progress")
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        agent_status = st.empty()
        
        # Agent status display
        agents = [
            ("📋", "Document Analyzer", "pending"),
            ("⚠️", "Risk Assessor", "pending"),
            ("🎯", "Negotiation Strategist", "pending"),
            ("⚖️", "Legal Advisor", "pending"),
            ("📊", "Market Researcher", "pending"),
            ("🔧", "Contract Optimizer", "pending"),
        ]
        
        def update_progress(progress):
            step = progress.current_step
            progress_bar.progress(step / 6)
            status_text.markdown(f"**{progress.current_agent}**: {progress.message}")
            
            # Update agent status
            updated_agents = []
            for i, (emoji, name, _) in enumerate(agents):
                if i + 1 < step:
                    updated_agents.append((emoji, name, "✅"))
                elif i + 1 == step:
                    updated_agents.append((emoji, name, "🔄"))
                else:
                    updated_agents.append((emoji, name, "⏳"))
            
            status_html = " → ".join([f"{e} {n} {s}" for e, n, s in updated_agents])
            agent_status.markdown(status_html)
        
        try:
            start_time = time.time()
            
            playbook = orchestrator.run_full_analysis(
                contract_text=st.session_state.document_text,
                document_name=st.session_state.document_name,
                context=st.session_state.negotiation_context,
                progress_callback=update_progress
            )
            
            elapsed = time.time() - start_time
            
            st.session_state.negotiation_playbook = playbook
            
            progress_bar.progress(1.0)
            status_text.success(f"✅ Analysis complete in {elapsed:.1f} seconds!")
            
            # Show agent timing
            timing = orchestrator.get_agent_timing()
            timing_text = " | ".join([f"{k}: {v:.1f}s" for k, v in timing.items()])
            st.caption(f"Agent timing: {timing_text}")
            
            time.sleep(1)
            st.rerun()
            
        except Exception as e:
            st.error(f"❌ Analysis failed: {str(e)}")
            import traceback
            st.code(traceback.format_exc())


def display_negotiation_playbook(playbook: 'NegotiationPlaybook'):
    """Display the complete negotiation playbook"""
    
    # Create tabs for different sections
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "📋 Executive Summary",
        "⚠️ Risk Analysis",
        "🎯 Negotiation Strategy", 
        "⚖️ Legal Review",
        "📊 Market Intel",
        "🗺️ Action Roadmap",
        "📝 Full Report"
    ])
    
    # Tab 1: Executive Summary
    with tab1:
        opt = playbook.optimization
        
        # Top-level metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            risk_color = {
                "CRITICAL": "#dc3545",
                "HIGH": "#fd7e14", 
                "MEDIUM": "#ffc107",
                "LOW": "#28a745"
            }.get(playbook.risk_assessment.overall_level, "#6c757d")
            
            st.markdown(f"""
            <div class="metric-card">
                <div style="font-size: 14px; color: #666;">Risk Score</div>
                <div style="font-size: 36px; font-weight: bold; color: {risk_color};">
                    {playbook.risk_assessment.overall_score}/100
                </div>
                <div style="font-size: 12px; color: {risk_color};">{playbook.risk_assessment.overall_level}</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            power = playbook.negotiation_strategy.power_balance
            power_color = "#28a745" if power > 0 else "#dc3545" if power < 0 else "#6c757d"
            power_text = "In Your Favor" if power > 0 else "Against You" if power < 0 else "Neutral"
            
            st.markdown(f"""
            <div class="metric-card">
                <div style="font-size: 14px; color: #666;">Power Balance</div>
                <div style="font-size: 36px; font-weight: bold; color: {power_color};">
                    {power:+.1f}
                </div>
                <div style="font-size: 12px; color: {power_color};">{power_text}</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            market_score = playbook.market_research.overall_favorability_score
            market_color = "#28a745" if market_score >= 60 else "#ffc107" if market_score >= 40 else "#dc3545"
            
            st.markdown(f"""
            <div class="metric-card">
                <div style="font-size: 14px; color: #666;">Market Position</div>
                <div style="font-size: 36px; font-weight: bold; color: {market_color};">
                    {market_score}/100
                </div>
                <div style="font-size: 12px; color: #666;">Favorability</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown(f"""
            <div class="metric-card">
                <div style="font-size: 14px; color: #666;">Success Rate</div>
                <div style="font-size: 36px; font-weight: bold; color: #2196F3;">
                    {opt.estimated_success_rate}
                </div>
                <div style="font-size: 12px; color: #666;">Estimated</div>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Assessment and Recommendation
        rec_color = {
            "DO NOT SIGN": "#dc3545",
            "PROCEED WITH CAUTION": "#fd7e14",
            "NEGOTIATE FIRST": "#ffc107",
            "ACCEPTABLE": "#28a745"
        }
        
        for key, color in rec_color.items():
            if key in opt.recommendation.upper():
                st.markdown(f"""
                <div style="background-color: {color}22; border-left: 4px solid {color}; 
                            padding: 15px; border-radius: 8px; margin: 10px 0;">
                    <strong style="color: {color};">📢 RECOMMENDATION: {opt.recommendation}</strong>
                </div>
                """, unsafe_allow_html=True)
                break
        else:
            st.info(f"📢 **Recommendation:** {opt.recommendation}")
        
        # Key Insights
        st.subheader("🔑 Key Insights")
        for insight in opt.key_insights:
            st.markdown(f"• {insight}")
        
        # Critical Decisions
        if opt.critical_decisions:
            st.subheader("⚠️ Critical Decisions Required")
            for decision in opt.critical_decisions:
                with st.expander(f"❓ {decision.decision}", expanded=True):
                    st.markdown(f"**Recommendation:** {decision.recommendation}")
                    st.markdown(f"**Rationale:** {decision.rationale}")
                    st.markdown(f"**Business Impact:** {decision.business_impact}")
                    st.markdown(f"**Decision Maker:** {decision.decision_maker}")
                    if decision.alternative:
                        st.markdown(f"**Alternative:** {decision.alternative}")
    
    # Tab 2: Risk Analysis
    with tab2:
        risk = playbook.risk_assessment
        
        st.markdown(f"### Overall Risk: {risk.overall_level} ({risk.overall_score}/100)")
        st.markdown(risk.summary)
        
        # Risk counts
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("🚨 Critical", risk.critical_count)
        col2.metric("🔶 High", risk.high_count)
        col3.metric("⚠️ Medium", risk.medium_count)
        col4.metric("✅ Low", risk.low_count)
        
        st.markdown("---")
        
        # Critical Risks
        if risk.critical_risks:
            st.subheader("🚨 Critical Risks")
            for r in risk.critical_risks:
                with st.expander(f"❌ {r.clause} (Score: {r.score}/100)", expanded=True):
                    st.markdown(f"**Category:** {r.category}")
                    st.markdown(f"**Description:** {r.description}")
                    st.markdown(f"**Impact:** {r.impact}")
                    st.markdown(f"**Likelihood:** {r.likelihood}")
                    st.markdown(f"**Financial Exposure:** {r.financial_exposure}")
                    if r.mitigation:
                        st.markdown("**Mitigation:**")
                        for k, v in r.mitigation.items():
                            st.markdown(f"  - {k}: {v}")
        
        # High Risks
        if risk.high_risks:
            st.subheader("🔶 High Priority Risks")
            for r in risk.high_risks[:5]:
                with st.expander(f"⚠️ {r.clause} (Score: {r.score}/100)"):
                    st.markdown(f"**Description:** {r.description}")
                    st.markdown(f"**Impact:** {r.impact}")
                    if r.mitigation:
                        st.markdown(f"**Action:** {r.mitigation.get('required_action', 'Review required')}")
    
    # Tab 3: Negotiation Strategy
    with tab3:
        strategy = playbook.negotiation_strategy
        
        # Power dynamics
        st.subheader("💪 Power Dynamics")
        st.markdown(strategy.power_interpretation)
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**✅ Factors in Your Favor:**")
            for f in strategy.factors_in_favor:
                st.markdown(f"• {f}")
        with col2:
            st.markdown("**❌ Factors Against You:**")
            for f in strategy.factors_against:
                st.markdown(f"• {f}")
        
        st.markdown("---")
        
        # BATNA
        col1, col2 = st.columns(2)
        with col1:
            st.info(f"**Your BATNA:** {strategy.your_batna}")
        with col2:
            st.info(f"**Their BATNA:** {strategy.their_batna}")
        
        st.markdown("---")
        
        # Priorities with talking points
        st.subheader("🎯 Negotiation Priorities")
        for p in strategy.priorities[:5]:
            with st.expander(f"#{p.rank}: {p.issue} (Leverage: {p.leverage_score}/10)", expanded=p.rank <= 2):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**Current Position:** {p.current_position}")
                    st.markdown(f"**Target Position:** {p.target_position}")
                    st.markdown(f"**Walk-Away Point:** {p.minimum_acceptable}")
                with col2:
                    st.markdown(f"**Strategy:** {p.strategy}")
                    st.markdown(f"**Rationale:** {p.rationale}")
                
                if p.talking_points:
                    st.markdown("**💬 Talking Points:**")
                    for tp in p.talking_points:
                        st.markdown(f"• \"{tp}\"")
                
                if p.counter_proposal:
                    st.markdown(f"**📝 Proposed Language:** {p.counter_proposal}")
        
        # Quick Wins
        if strategy.quick_wins:
            st.subheader("⚡ Quick Wins (High Success Probability)")
            for qw in strategy.quick_wins:
                with st.expander(f"✨ {qw.issue} - Likelihood: {qw.likelihood}"):
                    st.markdown(f"**Current:** {qw.current}")
                    st.markdown(f"**Request:** {qw.request}")
                    st.markdown(f"**Script:** \"{qw.script}\"")
        
        # Deal Breakers
        if strategy.deal_breakers:
            st.subheader("🚫 Deal Breakers")
            for db in strategy.deal_breakers:
                st.error(f"❌ {db}")
    
    # Tab 4: Legal Review
    with tab4:
        legal = playbook.legal_advisory
        
        st.markdown(f"### Legal Assessment: {legal.overall_assessment}")
        st.markdown(f"**Recommended Legal Review:** {'Yes' if legal.recommended_legal_review else 'No'}")
        
        col1, col2, col3 = st.columns(3)
        col1.metric("Compliance Issues", legal.compliance_issues_count)
        col2.metric("Enforceability Risks", legal.enforceability_risks_count)
        col3.metric("Major Concerns", legal.major_concerns_count)
        
        st.markdown("---")
        
        # Compliance Issues
        if legal.compliance_issues:
            st.subheader("📋 Compliance Issues")
            for c in legal.compliance_issues:
                severity_color = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢"}.get(c.severity, "⚪")
                with st.expander(f"{severity_color} {c.issue} ({c.jurisdiction})"):
                    st.markdown(f"**Requirement:** {c.requirement}")
                    st.markdown(f"**Contract Says:** {c.contract_provision}")
                    st.markdown(f"**Status:** {c.compliance_status}")
                    st.markdown(f"**Risk:** {c.risk}")
                    st.markdown(f"**Recommendation:** {c.recommendation}")
        
        # Enforceability Concerns
        if legal.enforceability_concerns:
            st.subheader("⚖️ Enforceability Concerns")
            for e in legal.enforceability_concerns:
                with st.expander(f"⚠️ {e.clause}"):
                    st.markdown(f"**Issue:** {e.issue}")
                    st.markdown(f"**Legal Principle:** {e.legal_principle}")
                    st.markdown(f"**Likelihood Struck Down:** {e.likelihood_struck_down}")
                    st.markdown(f"**Recommendation:** {e.recommendation}")
        
        # Missing Clauses
        if legal.missing_clauses:
            st.subheader("📝 Missing Standard Clauses")
            for clause in legal.missing_clauses:
                st.warning(f"Missing: {clause}")
    
    # Tab 5: Market Intel
    with tab5:
        market = playbook.market_research
        
        st.markdown(f"### Market Favorability: {market.overall_favorability_score}/100")
        st.markdown(market.overall_interpretation)
        
        st.info(f"**Industry:** {market.industry} | **Contract Type:** {market.contract_type}")
        st.info(f"**Typical Value Range:** {market.typical_contract_value}")
        
        st.markdown("---")
        
        # Benchmark Comparisons
        st.subheader("📊 Term Benchmarks vs Market")
        
        for b in market.benchmark_comparisons:
            assessment_colors = {
                "FAVORABLE": ("🟢", "#28a745"),
                "NEUTRAL": ("🟡", "#ffc107"),
                "UNFAVORABLE": ("🔴", "#dc3545"),
                "FAR_BELOW_MARKET": ("🔴🔴", "#dc3545"),
                "ABOVE_MARKET": ("🟢", "#28a745")
            }
            icon, color = assessment_colors.get(b.assessment, ("⚪", "#6c757d"))
            
            with st.expander(f"{icon} {b.term_category}: {b.assessment}"):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**This Contract:** {b.this_contract}")
                    st.markdown(f"**Percentile:** {b.percentile}")
                with col2:
                    st.markdown(f"**Market Standard:** {b.market_standard}")
                    st.markdown(f"**Impact:** {b.impact}")
                if b.recommendation:
                    st.markdown(f"**Recommendation:** {b.recommendation}")
        
        # Industry Trends
        if market.industry_trends:
            st.subheader("📈 Industry Trends")
            for trend in market.industry_trends:
                st.markdown(f"• {trend}")
    
    # Tab 6: Action Roadmap
    with tab6:
        opt = playbook.optimization
        
        st.markdown(f"### Recommended Timeline: {opt.recommended_timeline}")
        
        # Phase 1: Critical
        if opt.phase_1_critical:
            st.subheader("🔴 Phase 1: Critical Issues (Must Address)")
            for item in opt.phase_1_critical:
                with st.expander(f"#{item.rank} {item.issue} - {item.success_likelihood}", expanded=True):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**Current:** {item.current}")
                        st.markdown(f"**Target:** {item.target}")
                        st.markdown(f"**Minimum:** {item.minimum}")
                    with col2:
                        st.markdown(f"**Strategy:** {item.strategy}")
                        st.markdown(f"**If Rejected:** {item.if_rejected}")
                    
                    if item.talking_points:
                        st.markdown("**Talking Points:**")
                        for tp in item.talking_points:
                            st.markdown(f"• {tp}")
        
        # Phase 2: High Priority
        if opt.phase_2_high_priority:
            st.subheader("🟡 Phase 2: High Priority Issues")
            for item in opt.phase_2_high_priority:
                with st.expander(f"#{item.rank} {item.issue} - {item.success_likelihood}"):
                    st.markdown(f"**Current → Target:** {item.current} → {item.target}")
                    st.markdown(f"**Strategy:** {item.strategy}")
        
        # Phase 3: Optimization
        if opt.phase_3_optimization:
            st.subheader("🟢 Phase 3: Optimization")
            for item in opt.phase_3_optimization:
                with st.expander(f"#{item.rank} {item.issue}"):
                    st.markdown(f"**Current → Target:** {item.current} → {item.target}")
        
        # Next Steps
        st.markdown("---")
        st.subheader("📋 Next Steps")
        for i, step in enumerate(opt.next_steps, 1):
            st.markdown(f"{i}. {step}")
    
    # Tab 7: Full Report
    with tab7:
        st.markdown("### 📄 Full Negotiation Playbook")
        st.markdown("Download the complete report or view the executive summary below.")
        
        st.download_button(
            "📥 Download Full Report (Markdown)",
            data=playbook.executive_summary,
            file_name=f"negotiation_playbook_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
            mime="text/markdown"
        )
        
        st.markdown("---")
        st.markdown(playbook.executive_summary)


# ============================================
# LANGUAGE TRANSLATION
# ============================================

def run_translation():
    """Run the Language Translation page"""
    st.markdown("""
    <div class="nav-header">
        <h1>🌐 भाषा अनुवाद | Language Translation</h1>
        <p>Translate Legal Documents into 22+ Indian Languages</p>
    </div>
    """, unsafe_allow_html=True)
    
    if not TRANSLATION_AVAILABLE:
        st.error("❌ Translation module not available. Please check installation.")
        return
    
    if not GROQ_API_KEY:
        st.error("❌ Please set GROQ_API_KEY in your .env file")
        return
    
    # Initialize translator
    translator = LegalTranslator(api_key=GROQ_API_KEY)
    
    # Language selector with flags
    st.markdown("### 🗣️ Select Target Language")
    
    # Create language options
    language_options = []
    for code, details in SUPPORTED_LANGUAGES.items():
        display_name = f"{details['native_name']} ({details['name']})"
        language_options.append((code, display_name))
    
    # Sort by name
    language_options.sort(key=lambda x: x[1])
    
    # Create columns for popular languages
    st.markdown("**Popular Languages:**")
    pop_cols = st.columns(6)
    
    popular_langs = [
        ("hi", "🇮🇳 हिन्दी"),
        ("bn", "🇮🇳 বাংলা"),
        ("te", "🇮🇳 తెలుగు"),
        ("ta", "🇮🇳 தமிழ்"),
        ("mr", "🇮🇳 मराठी"),
        ("gu", "🇮🇳 ગુજરાતી")
    ]
    
    for i, (code, label) in enumerate(popular_langs):
        with pop_cols[i]:
            if st.button(label, key=f"quick_lang_{code}", use_container_width=True):
                st.session_state.selected_language = code
    
    # Full language dropdown
    col1, col2 = st.columns([2, 1])
    with col1:
        selected_idx = 0
        for i, (code, _) in enumerate(language_options):
            if code == st.session_state.selected_language:
                selected_idx = i
                break
        
        selected = st.selectbox(
            "All Languages",
            options=language_options,
            index=selected_idx,
            format_func=lambda x: x[1],
            key="lang_select"
        )
        st.session_state.selected_language = selected[0]
    
    with col2:
        target_lang = SUPPORTED_LANGUAGES.get(st.session_state.selected_language, {})
        st.info(f"**Script:** {target_lang.get('script', 'N/A')}\n\n**Region:** {target_lang.get('region', 'N/A')}")
    
    st.markdown("---")
    
    # Translation options
    tab1, tab2, tab3, tab4 = st.tabs([
        "📄 Translate Document",
        "📝 Translate Text",
        "📊 Translate Analysis",
        "📚 Translation History"
    ])
    
    # Tab 1: Translate Document
    with tab1:
        st.markdown("### Translate Uploaded Document")
        
        if not st.session_state.document_loaded:
            st.warning("📄 Please upload a document using the sidebar first.")
        else:
            st.success(f"📄 Document: **{st.session_state.document_name}**")
            
            # Show preview
            with st.expander("📖 Document Preview (First 1000 chars)"):
                st.text(st.session_state.document_text[:1000] + "...")
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                translate_full = st.button(
                    f"🌐 Translate Full Document to {target_lang.get('native_name', 'Selected Language')}",
                    type="primary",
                    use_container_width=True
                )
            
            with col2:
                translate_summary = st.button(
                    "📋 Translate Summary Only",
                    use_container_width=True
                )
            
            if translate_full:
                with st.spinner(f"Translating to {target_lang.get('name')}..."):
                    # Translate in chunks if document is large
                    text = st.session_state.document_text
                    
                    if len(text) > 5000:
                        st.info("📝 Large document detected. Translating in sections...")
                        
                        # Split into paragraphs/sections
                        sections = text.split('\n\n')
                        translated_sections = []
                        
                        progress = st.progress(0)
                        for i, section in enumerate(sections[:20]):  # Limit to 20 sections
                            if section.strip():
                                result = translator.translate(
                                    text=section,
                                    target_language=st.session_state.selected_language,
                                    context="legal document section"
                                )
                                translated_sections.append(result.translated_text)
                            progress.progress((i + 1) / min(len(sections), 20))
                        
                        full_translation = "\n\n".join(translated_sections)
                    else:
                        result = translator.translate(
                            text=text,
                            target_language=st.session_state.selected_language,
                            context="legal document"
                        )
                        full_translation = result.translated_text
                    
                    # Store translation
                    st.session_state.translated_content['full_document'] = {
                        'text': full_translation,
                        'language': st.session_state.selected_language,
                        'language_name': target_lang.get('name'),
                        'native_name': target_lang.get('native_name')
                    }
                    
                    # Add to history
                    st.session_state.translation_history.append({
                        'type': 'Document',
                        'language': target_lang.get('name'),
                        'preview': full_translation[:100] + "..."
                    })
                
                st.success(f"✅ Document translated to {target_lang.get('native_name')}!")
            
            # Display translated document if available
            if 'full_document' in st.session_state.translated_content:
                trans = st.session_state.translated_content['full_document']
                st.markdown(f"### 📄 Translated Document ({trans['native_name']})")
                
                st.markdown(f"""
                <div style="background-color: #f0f7ff; padding: 20px; border-radius: 10px; 
                            border-left: 4px solid #2196F3; font-size: 16px; line-height: 1.8;
                            color: #1a1a1a;">
                    {trans['text'][:3000]}{'...' if len(trans['text']) > 3000 else ''}
                </div>
                """, unsafe_allow_html=True)
                
                # Download button
                st.download_button(
                    f"📥 Download Translation ({trans['native_name']})",
                    data=trans['text'],
                    file_name=f"translated_{st.session_state.document_name}_{trans['language']}.txt",
                    mime="text/plain"
                )
    
    # Tab 2: Translate Custom Text
    with tab2:
        st.markdown("### Translate Custom Text")
        
        input_text = st.text_area(
            "Enter text to translate",
            height=200,
            placeholder="Paste any legal text here...\n\nFor example:\n'The indemnifying party shall hold harmless the indemnified party from any claims arising from...'"
        )
        
        if st.button("🌐 Translate Text", type="primary", disabled=not input_text):
            with st.spinner("Translating..."):
                result = translator.translate(
                    text=input_text,
                    target_language=st.session_state.selected_language,
                    context="legal text"
                )
                
                st.session_state.translated_content['custom_text'] = {
                    'original': input_text,
                    'translated': result.translated_text,
                    'language': target_lang.get('name'),
                    'native_name': target_lang.get('native_name')
                }
        
        if 'custom_text' in st.session_state.translated_content:
            trans = st.session_state.translated_content['custom_text']
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Original (English):**")
                st.markdown(f"""
                <div style="background-color: #fff3e0; padding: 15px; border-radius: 8px; 
                            border-left: 4px solid #ff9800; min-height: 150px; color: #1a1a1a;">
                    {trans['original']}
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"**Translated ({trans['native_name']}):**")
                st.markdown(f"""
                <div style="background-color: #e8f5e9; padding: 15px; border-radius: 8px; 
                            border-left: 4px solid #4caf50; min-height: 150px; color: #1a1a1a;">
                    {trans['translated']}
                </div>
                """, unsafe_allow_html=True)
    
    # Tab 3: Translate Analysis Results
    with tab3:
        st.markdown("### Translate Analysis Results")
        
        if not st.session_state.analysis_result and not st.session_state.negotiation_playbook:
            st.info("📊 Run a Risk Assessment or NegotiateAI analysis first, then translate the results here.")
        else:
            # Risk Assessment translation
            if st.session_state.analysis_result:
                st.markdown("#### 📊 Risk Assessment Results")
                
                if st.button(f"🌐 Translate Risk Summary to {target_lang.get('native_name')}", key="trans_risk"):
                    with st.spinner("Translating risk assessment..."):
                        # Get summary from analysis
                        analysis = st.session_state.analysis_result
                        
                        summary_text = f"""
Risk Assessment Summary:
Overall Risk Score: {analysis.overall_score}/100
Risk Level: {analysis.severity_level.value if hasattr(analysis.severity_level, 'value') else analysis.severity_level}
Total Risky Clauses: {len(analysis.clause_risks)}

Key Findings:
"""
                        for i, risk in enumerate(analysis.clause_risks[:5], 1):
                            summary_text += f"{i}. {risk.clause_text[:100]}... - {risk.risk_category.value if hasattr(risk.risk_category, 'value') else risk.risk_category}\n"
                        
                        result = translator.translate(
                            text=summary_text,
                            target_language=st.session_state.selected_language,
                            context="legal risk assessment"
                        )
                        
                        st.session_state.translated_content['risk_summary'] = {
                            'text': result.translated_text,
                            'language': target_lang.get('name'),
                            'native_name': target_lang.get('native_name')
                        }
                
                if 'risk_summary' in st.session_state.translated_content:
                    trans = st.session_state.translated_content['risk_summary']
                    st.markdown(f"**Risk Summary in {trans['native_name']}:**")
                    st.markdown(f"""
                    <div style="background-color: #ffebee; padding: 15px; border-radius: 8px; 
                                border-left: 4px solid #f44336; color: #1a1a1a;">
                        {trans['text']}
                    </div>
                    """, unsafe_allow_html=True)
            
            # NegotiateAI translation
            if st.session_state.negotiation_playbook:
                st.markdown("#### 🤝 NegotiateAI Results")
                
                if st.button(f"🌐 Translate Negotiation Strategy to {target_lang.get('native_name')}", key="trans_negotiate"):
                    with st.spinner("Translating negotiation strategy..."):
                        playbook = st.session_state.negotiation_playbook
                        
                        strategy_text = f"""
Negotiation Strategy Summary:

Overall Assessment: {playbook.optimization.overall_assessment}
Recommendation: {playbook.optimization.recommendation}
Success Rate: {playbook.optimization.estimated_success_rate}

Key Insights:
"""
                        for insight in playbook.optimization.key_insights[:5]:
                            strategy_text += f"• {insight}\n"
                        
                        strategy_text += "\nDeal Breakers:\n"
                        for db in playbook.negotiation_strategy.deal_breakers[:3]:
                            strategy_text += f"• {db}\n"
                        
                        result = translator.translate(
                            text=strategy_text,
                            target_language=st.session_state.selected_language,
                            context="negotiation strategy"
                        )
                        
                        st.session_state.translated_content['negotiation_summary'] = {
                            'text': result.translated_text,
                            'language': target_lang.get('name'),
                            'native_name': target_lang.get('native_name')
                        }
                
                if 'negotiation_summary' in st.session_state.translated_content:
                    trans = st.session_state.translated_content['negotiation_summary']
                    st.markdown(f"**Negotiation Strategy in {trans['native_name']}:**")
                    st.markdown(f"""
                    <div style="background-color: #e3f2fd; padding: 15px; border-radius: 8px; 
                                border-left: 4px solid #2196f3; color: #1a1a1a;">
                        {trans['text']}
                    </div>
                    """, unsafe_allow_html=True)
    
    # Tab 4: Translation History
    with tab4:
        st.markdown("### 📚 Translation History")
        
        if not st.session_state.translation_history:
            st.info("No translations yet. Translate some content to see history here.")
        else:
            for i, item in enumerate(reversed(st.session_state.translation_history[-10:])):
                with st.expander(f"{item['type']} → {item['language']}"):
                    st.markdown(item['preview'])
            
            if st.button("🗑️ Clear History"):
                st.session_state.translation_history = []
                st.session_state.translated_content = {}
                st.rerun()
    
    # Language info footer
    st.markdown("---")
    st.markdown("### 🇮🇳 Supported Indian Languages")
    
    # Display all languages in a grid
    lang_cols = st.columns(4)
    languages_list = list(SUPPORTED_LANGUAGES.items())
    
    for i, (code, details) in enumerate(languages_list):
        with lang_cols[i % 4]:
            st.markdown(f"**{details['native_name']}** ({details['name']})")


# ============================================
# HOME PAGE
# ============================================

def run_home():
    """Run the Home page"""
    st.markdown("""
    <div class="nav-header">
        <h1>⚖️ ClauseCare</h1>
        <p>Your Intelligent Legal Document Platform</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    Welcome to ClauseCare - a comprehensive platform for analyzing and understanding legal documents.
    Upload a document to get started!
    """)
    
    # Feature cards - Row 1
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        ### 🔍 Risk Assessment
        - Identify risky clauses
        - Get severity scores (0-100)
        - AI-powered deep analysis
        - Actionable recommendations
        - Export reports
        """)
        if st.button("Go to Risk Assessment", key="nav_risk", use_container_width=True):
            st.session_state.current_page = "risk"
            st.rerun()
    
    with col2:
        st.markdown("""
        ### 🤝 NegotiateAI
        - 6-agent AI analysis team
        - Negotiation strategies & tactics
        - Market benchmarking
        - Legal compliance review
        - Complete playbook generation
        """)
        if st.button("Go to NegotiateAI", key="nav_negotiate", use_container_width=True):
            st.session_state.current_page = "negotiate"
            st.rerun()
    
    # Feature cards - Row 2
    col3, col4 = st.columns(2)
    
    with col3:
        st.markdown("""
        ### 💬 Document Chat
        - Ask questions in plain language
        - Get cited answers
        - Source references
        - Legal term explanations
        - Conversation history
        """)
        if st.button("Go to Document Chat", key="nav_chat", use_container_width=True):
            st.session_state.current_page = "chat"
            st.rerun()
    
    with col4:
        st.markdown("""
        ### ✨ Clause Simplifier
        - Plain English translations
        - Flesch readability scores
        - Before/after comparison
        - Batch simplification
        - Legal jargon removal
        """)
        if st.button("Go to Clause Simplifier", key="nav_simplify", use_container_width=True):
            st.session_state.current_page = "simplify"
            st.rerun()
    
    # Feature cards - Row 3
    col5, col6 = st.columns(2)
    
    with col5:
        st.markdown("""
        ### 🌐 Language Translation
        - 22+ Indian languages
        - Legal document translation
        - Translate analysis results
        - Regional language support
        - हिन्दी, বাংলা, తెలుగు, தமிழ்...
        """)
        if st.button("Go to Translation", key="nav_translate", use_container_width=True):
            st.session_state.current_page = "translate"
            st.rerun()
    
    with col6:
        st.markdown("""
        ### 📄 Document Processor
        - PDF, DOCX, TXT support
        - OCR for scanned docs
        - Table extraction
        - Readability analysis
        - Metadata extraction
        """)
        if st.button("Go to Document Processor", key="nav_processor", use_container_width=True):
            st.session_state.current_page = "processor"
            st.rerun()
    
    # Status
    st.markdown("---")
    st.subheader("📊 System Status")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if GROQ_API_KEY:
            st.success("✅ Groq API Connected")
        else:
            st.error("❌ Groq API Key Missing")
        
        if RISK_ASSESSMENT_AVAILABLE:
            st.success("✅ Risk Assessment Ready")
        else:
            st.error("❌ Risk Assessment Unavailable")
    
    with col2:
        if RAG_CHATBOT_AVAILABLE:
            st.success("✅ RAG Chatbot Ready")
        else:
            st.error("❌ RAG Chatbot Unavailable")
        
        if NEGOTIATE_AI_AVAILABLE:
            st.success("✅ NegotiateAI Ready")
        else:
            st.warning("⚠️ NegotiateAI Unavailable")
    
    with col3:
        if ADVANCED_PROCESSOR_AVAILABLE:
            st.success("✅ Advanced Processor Ready")
        else:
            st.warning("⚠️ Using Basic Processor")
        
        if CLAUSE_SIMPLIFICATION_AVAILABLE:
            st.success("✅ Clause Simplifier Ready")
        else:
            st.warning("⚠️ Simplifier Limited")
    
    with col4:
        if TRANSLATION_AVAILABLE:
            st.success("✅ Translation Ready")
            st.caption("22+ Indian Languages")
        else:
            st.warning("⚠️ Translation Unavailable")
        
        if AUTH_AVAILABLE and DATABASE_AVAILABLE:
            st.success("✅ Auth & History Ready")
        else:
            st.warning("⚠️ Auth Unavailable")


# ============================================
# HISTORY PAGE
# ============================================

def run_history():
    """Run the analysis history page"""
    st.markdown("""
    <div class="nav-header">
        <h1>📊 Analysis History</h1>
        <p>View your saved document analyses</p>
    </div>
    """, unsafe_allow_html=True)
    
    if not AUTH_AVAILABLE or not DATABASE_AVAILABLE:
        st.error("History feature requires authentication. Please check installation.")
        return
    
    auth = AuthManager()
    
    if not auth.is_authenticated():
        st.warning("🔒 Please login to view your analysis history")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### Login")
            login_form()
        with col2:
            st.markdown("### Register")
            register_form()
        return
    
    username = auth.get_current_user()
    db = get_database()
    
    # User stats
    stats = db.get_user_stats(username)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Analyses", stats.get("total_analyses", 0))
    with col2:
        st.metric("Documents Analyzed", stats.get("documents_analyzed", 0))
    with col3:
        st.metric("Risk Analyses", stats.get("by_type", {}).get("risk", 0))
    with col4:
        st.metric("Negotiations", stats.get("by_type", {}).get("negotiate", 0))
    
    st.markdown("---")
    
    # Filter options
    col1, col2 = st.columns([2, 1])
    with col1:
        filter_type = st.selectbox(
            "Filter by type",
            ["All", "risk", "negotiate", "chat", "translate", "simplify"],
            format_func=lambda x: {
                "All": "📋 All Analyses",
                "risk": "🔍 Risk Assessments",
                "negotiate": "🤝 Negotiations",
                "chat": "💬 Chat Sessions",
                "translate": "🌐 Translations",
                "simplify": "✨ Simplifications"
            }.get(x, x)
        )
    with col2:
        if st.button("🗑️ Clear All History", type="secondary"):
            if st.session_state.get("confirm_clear_history"):
                deleted = db.clear_user_history(username)
                st.success(f"Cleared {deleted} entries")
                st.session_state.confirm_clear_history = False
                st.rerun()
            else:
                st.session_state.confirm_clear_history = True
                st.warning("Click again to confirm deletion")
    
    # Get history
    analysis_type = None if filter_type == "All" else filter_type
    history = db.get_user_history(username, analysis_type=analysis_type, limit=50)
    
    if not history:
        st.info("No analysis history found. Start analyzing documents to build your history!")
        return
    
    # Display history
    for entry in history:
        type_emoji = {
            "risk": "🔍",
            "negotiate": "🤝",
            "chat": "💬",
            "translate": "🌐",
            "simplify": "✨"
        }.get(entry.analysis_type, "📄")
        
        with st.expander(f"{type_emoji} {entry.document_name} - {entry.analysis_type.title()} ({entry.created_at[:10]})"):
            st.markdown(entry.summary)
            
            col1, col2 = st.columns([3, 1])
            with col2:
                if st.button("🗑️ Delete", key=f"del_{entry.id}"):
                    if db.delete_analysis(entry.id, username):
                        st.success("Deleted")
                        st.rerun()


# ============================================
# MAIN APPLICATION
# ============================================

def main():
    """Main application entry point"""
    init_session_state()
    
    # Initialize auth manager if available
    auth = None
    if AUTH_AVAILABLE:
        auth = AuthManager()
    
    # Sidebar
    with st.sidebar:
        st.markdown("""
        <div style="text-align: center; padding: 1rem;">
            <h1>⚖️ ClauseCare</h1>
        </div>
        """, unsafe_allow_html=True)
        
        # User authentication section
        if AUTH_AVAILABLE and auth:
            if auth.is_authenticated():
                user = auth.get_current_user()
                st.markdown(f"""
                <div class="user-badge">
                    👤 <strong>{user}</strong>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns([1, 1])
                with col1:
                    if st.button("📊 History", use_container_width=True):
                        st.session_state.current_page = "history"
                        st.rerun()
                with col2:
                    if st.button("🚪 Logout", use_container_width=True):
                        auth.logout()
                        st.rerun()
            else:
                st.info("👤 Guest Mode")
                st.caption("Login to save analysis history")
                
                with st.expander("🔐 Login / Register"):
                    auth_tab = st.radio("Select Action", ["Login", "Register"], horizontal=True, label_visibility="collapsed")
                    
                    if auth_tab == "Login":
                        with st.form("sidebar_login"):
                            username = st.text_input("Username", placeholder="Enter username")
                            password = st.text_input("Password", type="password", placeholder="Enter password")
                            if st.form_submit_button("Login", use_container_width=True):
                                success, msg = auth.login(username, password)
                                if success:
                                    st.success(msg)
                                    st.rerun()
                                else:
                                    st.error(msg)
                    else:
                        with st.form("sidebar_register"):
                            username = st.text_input("Username", placeholder="Choose username")
                            email = st.text_input("Email", placeholder="Your email")
                            password = st.text_input("Password", type="password", placeholder="Min 6 chars")
                            confirm = st.text_input("Confirm", type="password", placeholder="Confirm password")
                            if st.form_submit_button("Register", use_container_width=True):
                                success, msg = auth.register(username, email, password, confirm)
                                if success:
                                    st.success(f"✅ {msg}! Please login.")
                                else:
                                    st.error(msg)
        
        st.markdown("---")
        
        # Navigation
        st.subheader("🧭 Navigation")
        
        if st.button("🏠 Home", use_container_width=True):
            st.session_state.current_page = "home"
            st.rerun()
        
        if st.button("🔍 Risk Assessment", use_container_width=True):
            st.session_state.current_page = "risk"
            st.rerun()
        
        if st.button("🤝 NegotiateAI", use_container_width=True):
            st.session_state.current_page = "negotiate"
            st.rerun()
        
        if st.button("💬 Document Chat", use_container_width=True):
            st.session_state.current_page = "chat"
            st.rerun()
        
        if st.button("✨ Clause Simplifier", use_container_width=True):
            st.session_state.current_page = "simplify"
            st.rerun()
        
        if st.button("🌐 Translation", use_container_width=True):
            st.session_state.current_page = "translate"
            st.rerun()
        
        if st.button("📄 Document Processor", use_container_width=True):
            st.session_state.current_page = "processor"
            st.rerun()
        
        st.markdown("---")
        
        # Document upload
        st.subheader("📤 Upload Document")
        
        uploaded_file = st.file_uploader(
            "Choose a legal document",
            type=["pdf", "docx", "txt"],
            help="Upload a contract or legal document"
        )
        
        if uploaded_file:
            # Check if it's a new file
            if st.session_state.document_name != uploaded_file.name:
                with st.spinner(f"Processing {uploaded_file.name}..."):
                    try:
                        text, metadata = extract_text_from_file(uploaded_file)
                        
                        if text and len(text.strip()) > 50:
                            st.session_state.document_text = text
                            st.session_state.document_name = uploaded_file.name
                            st.session_state.document_metadata = metadata
                            st.session_state.document_loaded = True
                            
                            # Reset analysis states for new document
                            st.session_state.analysis_result = None
                            st.session_state.conversation = None
                            st.session_state.chat_messages = []
                            
                            st.success(f"✅ Loaded {len(text):,} characters")
                        else:
                            st.error("Could not extract text from document.")
                            
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
        
        # Current document info
        if st.session_state.document_loaded:
            st.markdown("---")
            st.markdown(f"**Current Document:**")
            st.info(st.session_state.document_name)
            
            if st.button("🗑️ Clear Document", use_container_width=True):
                st.session_state.document_text = None
                st.session_state.document_name = None
                st.session_state.document_metadata = {}
                st.session_state.document_loaded = False
                st.session_state.analysis_result = None
                st.session_state.conversation = None
                st.session_state.chat_messages = []
                if st.session_state.vector_store:
                    st.session_state.vector_store.clear()
                st.rerun()
        
        st.markdown("---")
        
        # API status
        if GROQ_API_KEY:
            st.success("🔑 API Key Loaded")
        else:
            st.warning("⚠️ Set GROQ_API_KEY in .env")
    
    # Main content based on navigation
    page = st.session_state.current_page
    
    if page == "home":
        run_home()
    elif page == "risk":
        run_risk_assessment()
    elif page == "negotiate":
        run_negotiate_ai()
    elif page == "chat":
        run_document_chat()
    elif page == "processor":
        run_document_processor()
    elif page == "simplify":
        run_clause_simplification()
    elif page == "translate":
        run_translation()
    elif page == "history":
        run_history()
    else:
        run_home()


if __name__ == "__main__":
    main()
